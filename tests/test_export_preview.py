from __future__ import annotations

from io import BytesIO
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import uuid4

import docx
import pytest
from fastapi.testclient import TestClient
from pypdf import PdfReader


@pytest.fixture()
def client():
    from services.settings import load_settings as _ls
    _ls.cache_clear()

    fake_settings = MagicMock()
    fake_settings.supabase_jwt_secret = "test-jwt"
    fake_settings.allowed_origin = "http://localhost:3000"
    fake_settings.supabase_url = "https://fake.supabase.co"
    fake_settings.supabase_service_role_key = "fake-service-key"
    fake_settings.normalized_supabase_url = "https://fake.supabase.co"

    with patch("services.settings.load_settings", return_value=fake_settings):
        from main import create_app
        app = create_app()
        from routers.auth import get_supabase_admin, require_authenticated_user
        from routers.reports import get_report_repository
        from services.supabase_admin import UserProfile
        fake_supabase = MagicMock()
        fake_repository = MagicMock()
        fake_repository._supabase = fake_supabase
        app.dependency_overrides[get_supabase_admin] = lambda: fake_supabase
        app.dependency_overrides[get_report_repository] = lambda: fake_repository
        app.dependency_overrides[require_authenticated_user] = lambda: UserProfile(
            id=str(uuid4()), role="lanh_dao", village_id=None, force_password_reset=False
        )
        yield TestClient(app)
        app.dependency_overrides.clear()


def test_export_preview_empty_data(client: TestClient) -> None:
    # Get the fake_supabase mock
    fake_supabase = None
    for dep, stub_func in client.app.dependency_overrides.items():
        if "get_supabase_admin" in str(dep):
            fake_supabase = stub_func()
            break
    if fake_supabase is None:
        fake_supabase = MagicMock()

    # Stub Supabase REST response with no reports
    async def mock_rest_request(method, path, body=None, prefer=None):
        if "report_periods" in path:
            return [{"id": str(uuid4()), "name": "Quý III/2026", "due_date": "2026-09-30"}]
        elif "villages" in path:
            return [{"id": str(uuid4()), "name": "Thôn Tà Lang"}]
        elif "reports" in path:
            return [] # No reports submitted yet
        return []

    fake_supabase._rest_request = AsyncMock(side_effect=mock_rest_request)

    # 1. Test preview HTML
    response = client.get("/reports/preview/pdf?period_id=Quý III/2026")
    assert response.status_code == 200
    assert "text/html" in response.headers["content-type"]
    assert "Chưa có dữ liệu" in response.text

    # 2. Test export XLSX
    response = client.get("/reports/export/xlsx?period_id=Quý III/2026")
    assert response.status_code == 200
    assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    # ZIP signature check for xlsx
    assert response.content.startswith(b"PK")

    # 3. Test export DOCX
    response = client.get("/reports/export/docx?period_id=Quý III/2026")
    assert response.status_code == 200
    assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert response.content.startswith(b"PK")

    # 4. Test export PDF
    response = client.get("/reports/export/pdf?period_id=Quý III/2026")
    assert response.status_code == 200
    assert response.headers["content-type"] == "application/pdf"
    assert response.content.startswith(b"%PDF")


def test_export_preview_with_data(client: TestClient) -> None:
    # Get the fake_supabase mock
    fake_supabase = None
    for dep, stub_func in client.app.dependency_overrides.items():
        if "get_supabase_admin" in str(dep):
            fake_supabase = stub_func()
            break
    if fake_supabase is None:
        fake_supabase = MagicMock()

    period_uuid = str(uuid4())
    village_uuid = str(uuid4())
    report_uuid = str(uuid4())

    # Stub Supabase REST response with reports & values
    async def mock_rest_request(method, path, body=None, prefer=None):
        if "report_periods" in path:
            return [{"id": period_uuid, "name": "Quý III/2026", "due_date": "2026-09-30"}]
        elif "villages" in path:
            return [{"id": village_uuid, "name": "Thôn Tà Lang"}]
        elif "reports" in path:
            return [{
                "id": report_uuid,
                "village_id": village_uuid,
                "reporter_name": "Nguyen Van A",
                "reporter_phone": "0905123456",
                "workflow_status": "approved",
                "timeliness_status": "on_time",
                "submitted_at": "2026-07-07T12:00:00Z"
            }]
        elif "report_values" in path:
            return [
                {"report_id": report_uuid, "ct_code": "CT01", "value": 150},
                {"report_id": report_uuid, "ct_code": "CT02", "value": 520},
                {"report_id": report_uuid, "ct_code": "CT03", "value": 5},
                {"report_id": report_uuid, "ct_code": "CT04", "value": 10}
            ]
        return []

    fake_supabase._rest_request = AsyncMock(side_effect=mock_rest_request)

    # 1. Test HTML preview has report details
    response = client.get("/reports/preview/pdf?period_id=Quý III/2026")
    assert response.status_code == 200
    assert "Thôn Tà Lang" in response.text
    assert "150" in response.text

    # 2. Test export XLSX
    response = client.get("/reports/export/xlsx?period_id=Quý III/2026")
    assert response.status_code == 200
    assert response.content.startswith(b"PK")

    # 3. Test export DOCX
    response = client.get("/reports/export/docx?period_id=Quý III/2026")
    assert response.status_code == 200
    assert response.content.startswith(b"PK")

    # 4. Test export PDF
    response = client.get("/reports/export/pdf?period_id=Quý III/2026")
    assert response.status_code == 200
    assert response.content.startswith(b"%PDF")


def test_period_export_scope_uses_only_assigned_villages() -> None:
    from routers.reports import scope_villages_map_to_period

    villages_map = {
        "village-1": "Thôn An Sơn",
        "village-2": "Thôn Hòa Nhơn",
        "village-3": "Thôn Tà Lang",
    }

    assert scope_villages_map_to_period(
        villages_map,
        ["village-3", "village-1"],
    ) == {
        "village-3": "Thôn Tà Lang",
        "village-1": "Thôn An Sơn",
    }
    assert scope_villages_map_to_period(villages_map, []) == villages_map


def test_village_docx_export_contains_only_authorized_village(client: TestClient) -> None:
    fake_supabase = None
    for dependency, stub_func in client.app.dependency_overrides.items():
        if "get_supabase_admin" in str(dependency):
            fake_supabase = stub_func()
            break
    assert fake_supabase is not None

    period_uuid = str(uuid4())
    selected_village_uuid = str(uuid4())
    other_village_uuid = str(uuid4())
    selected_report_uuid = str(uuid4())
    other_report_uuid = str(uuid4())

    async def mock_rest_request(method, path, body=None, prefer=None):
        if "report_periods" in path:
            return [{"id": period_uuid, "name": "Tháng 7/2026"}]
        if "villages" in path:
            return [
                {"id": selected_village_uuid, "name": "Thôn An Sơn"},
                {"id": other_village_uuid, "name": "Thôn Hòa Nhơn"},
            ]
        if "/rest/v1/reports?" in path:
            return [
                {
                    "id": selected_report_uuid,
                    "village_id": selected_village_uuid,
                    "workflow_status": "approved",
                    "timeliness_status": "on_time",
                    "submitted_at": "2026-07-15T10:00:00Z",
                },
                {
                    "id": other_report_uuid,
                    "village_id": other_village_uuid,
                    "workflow_status": "approved",
                    "timeliness_status": "on_time",
                    "submitted_at": "2026-07-15T10:00:00Z",
                },
            ]
        if "/rest/v1/report_values?" in path:
            return [
                {"report_id": selected_report_uuid, "ct_code": "CT01", "value": 318},
                {"report_id": other_report_uuid, "ct_code": "CT01", "value": 421},
            ]
        return []

    fake_supabase._rest_request = AsyncMock(side_effect=mock_rest_request)

    response = client.get(
        f"/reports/village/{selected_village_uuid}/export/docx",
        params={"period_id": period_uuid},
    )

    assert response.status_code == 200
    assert response.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "Phieu_bao_cao_Th%C3%B4n_An_S%C6%A1n_Th%C3%A1ng_7%2F2026.docx" in (
        response.headers["content-disposition"]
    )
    document = docx.Document(BytesIO(response.content))
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    cell_text = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    assert "Phạm vi: Thôn An Sơn" in paragraph_text
    assert "Thôn An Sơn" in cell_text
    assert "318" in cell_text
    assert "Thôn Hòa Nhơn" not in cell_text
    assert "421" not in cell_text

    pdf_response = client.get(
        f"/reports/village/{selected_village_uuid}/export/pdf",
        params={"period_id": period_uuid},
    )
    assert pdf_response.status_code == 200
    assert pdf_response.headers["content-type"] == "application/pdf"
    assert "Phieu_bao_cao_Th%C3%B4n_An_S%C6%A1n_Th%C3%A1ng_7%2F2026.pdf" in (
        pdf_response.headers["content-disposition"]
    )
    pdf_text = "\n".join(
        page.extract_text() or ""
        for page in PdfReader(BytesIO(pdf_response.content)).pages
    )
    assert "Thôn An Sơn" in pdf_text
    assert "318" in pdf_text
    assert "Thôn Hòa Nhơn" not in pdf_text
    assert "421" not in pdf_text


@pytest.mark.asyncio
async def test_period_export_query_keeps_previously_submitted_needs_revision_report() -> None:
    from routers.reports import get_period_reports_data

    report_uuid = str(uuid4())
    village_uuid = str(uuid4())
    requested_paths: list[str] = []

    async def mock_rest_request(method, path, body=None, prefer=None):
        requested_paths.append(path)
        if "/rest/v1/reports?" in path:
            return [
                {
                    "id": report_uuid,
                    "village_id": village_uuid,
                    "workflow_status": "needs_revision",
                    "timeliness_status": "on_time",
                    "submitted_at": "2026-07-15T10:00:00Z",
                }
            ]
        if "/rest/v1/report_values?" in path:
            return [{"report_id": report_uuid, "ct_code": "CT01", "value": 320}]
        return []

    fake_supabase = MagicMock()
    fake_supabase._rest_request = AsyncMock(side_effect=mock_rest_request)

    result = await get_period_reports_data(fake_supabase, str(uuid4()))

    report_path = next(path for path in requested_paths if "/rest/v1/reports?" in path)
    assert "timeliness_status=in.(on_time,late)" in report_path
    assert "workflow_status=in." not in report_path
    assert "id,village_id,period_id,workflow_status" in report_path
    assert "submitted_at,approved_at" in report_path
    assert result[0]["workflow_status"] == "needs_revision"
    assert result[0]["values"]["CT01"] == 320
