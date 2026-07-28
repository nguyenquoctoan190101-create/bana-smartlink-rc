from __future__ import annotations

from io import BytesIO
from zipfile import ZipFile

import docx
import openpyxl
from docx.shared import Mm, Pt
from pypdf import PdfReader

from routers.reports import generate_docx_file, generate_pdf_file, generate_preview_html
from services.export_service import generate_summary_xlsx_file, generate_village_xlsx_file


def _report() -> tuple[list[dict], dict[str, str]]:
    return (
        [
            {
                "id": "report-1",
                "village_id": "village-1",
                "workflow_status": "approved",
                "timeliness_status": "on_time",
                "report_source": "excel",
                "version": 3,
                "rule_version": "2026-07",
                "submitted_at": "2026-07-13T10:00:00Z",
                "validation_flags": [{
                    "ct_code": "CT03",
                    "error_type": "TEXT",
                    "message": "=unsafe warning",
                    "resolved": False,
                }],
                "values": {
                    "CT01": 100,
                    "CT02": None,
                    "CT03": "=SUM(1,1)",
                    "CT14": 1,
                },
            }
        ],
        {"village-1": "<script>Village</script>"},
    )


def test_all_export_formats_keep_ct01_to_ct14_nulls_and_formula_safety() -> None:
    reports, villages = _report()
    period = "=Quarter 3/2026"

    summary = openpyxl.load_workbook(
        BytesIO(generate_summary_xlsx_file(period, reports, villages))
    )
    summary_sheet = summary["Bảng tổng hợp"]
    assert summary.sheetnames == [
        "Bảng tổng hợp",
        "Dashboard",
        "Theo dõi tiến độ",
        "Từ điển dữ liệu",
        "Cảnh báo dữ liệu",
        "Nguồn dữ liệu",
    ]
    assert "CT14" in str(summary_sheet["P4"].value)
    assert summary_sheet["D5"].value is None  # CT02 is intentionally blank.
    assert summary_sheet["F6"].value in (None, "")  # CT04 total stays incomplete.
    assert summary_sheet["E5"].value == "'=SUM(1,1)"
    assert not str(summary_sheet["A2"].value).startswith("=")
    assert summary["Dashboard"]["B4"].value == 1
    assert summary["Dashboard"]["B16"].value == "—"
    assert summary["Từ điển dữ liệu"]["A17"].value == "CT14"
    assert (
        summary["Từ điển dữ liệu"]["B8"].value
        == "Số người có công với cách mạng đang được quản lý"
    )
    assert summary["Cảnh báo dữ liệu"]["E2"].value == "'=unsafe warning"
    assert summary["Nguồn dữ liệu"]["C2"].value == "Excel"
    assert summary["Nguồn dữ liệu"]["D2"].value == 3
    assert summary["Nguồn dữ liệu"]["F2"].value == "Đã duyệt"
    assert summary["Nguồn dữ liệu"]["G2"].value == "Đúng hạn"

    village = openpyxl.load_workbook(
        BytesIO(generate_village_xlsx_file(period, reports[0], villages["village-1"]))
    )
    village_sheet = village["Phiếu báo cáo"]
    assert village_sheet["D15"].value is None
    assert village_sheet["D16"].value == "'=SUM(1,1)"
    assert village_sheet["D27"].value == 1  # CT14 row.

    document = docx.Document(BytesIO(generate_docx_file(period, reports, villages)))
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    assert "CT14" in cells
    assert "" in cells  # CT02 stays blank in the DOCX table.

    scoped_document = docx.Document(
        BytesIO(
            generate_docx_file(
                period,
                reports,
                villages,
                scope_name="Thôn An Sơn",
            )
        )
    )
    scoped_text = "\n".join(
        paragraph.text for paragraph in scoped_document.paragraphs
    )
    assert "Phạm vi: Thôn An Sơn" in scoped_text

    pdf = generate_pdf_file(period, reports, villages)
    assert pdf.startswith(b"%PDF-")
    assert len(pdf) > 1_000


def test_xlsx_exports_define_single_page_print_layouts() -> None:
    reports, villages = _report()
    summary = openpyxl.load_workbook(
        BytesIO(generate_summary_xlsx_file("Tháng 7/2026", reports, villages))
    )

    expected = {
        "Bảng tổng hợp": (8, "landscape", "$A$1:$Q$"),
        "Dashboard": (8, "landscape", "$A$1:$E$"),
        "Theo dõi tiến độ": (9, "landscape", "$A$1:$G$"),
        "Từ điển dữ liệu": (9, "landscape", "$A$1:$E$"),
        "Cảnh báo dữ liệu": (9, "landscape", "$A$1:$F$"),
        "Nguồn dữ liệu": (9, "landscape", "$A$1:$H$"),
    }
    for sheet_name, (paper_size, orientation, print_area) in expected.items():
        sheet = summary[sheet_name]
        assert sheet.sheet_properties.pageSetUpPr.fitToPage is True
        assert sheet.page_setup.fitToWidth == 1
        assert sheet.page_setup.fitToHeight == 1
        assert sheet.page_setup.paperSize == paper_size
        assert sheet.page_setup.orientation == orientation
        assert print_area in str(sheet.print_area)
        assert sheet.print_options.horizontalCentered is True
        assert "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in sheet.oddHeader.center.text
        assert sheet.oddFooter.center.text == "Trang &P/&N"
        assert sheet.page_margins.left == 1.18
        assert sheet.page_margins.right == 0.71
        assert sheet.sheet_view.showGridLines is False

    village = openpyxl.load_workbook(
        BytesIO(
            generate_village_xlsx_file(
                "Tháng 7/2026", reports[0], villages["village-1"]
            )
        )
    )["Phiếu báo cáo"]
    assert village.sheet_properties.pageSetUpPr.fitToPage is True
    assert village.page_setup.fitToWidth == 1
    assert village.page_setup.fitToHeight == 1
    assert village.page_setup.paperSize == 9
    assert village.page_setup.orientation == "portrait"
    assert "$A$1:$E$" in str(village.print_area)
    assert village.oddHeader.center.text is None
    assert village["A1"].value == "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
    assert village["A1"].font.name == "Times New Roman"
    assert village.oddFooter.center.text == "Trang &P/&N"


def test_docx_and_pdf_exports_use_readable_administrative_layout() -> None:
    reports, villages = _report()
    period = "Tháng 7/2026"

    document = docx.Document(
        BytesIO(generate_docx_file(period, reports, villages))
    )
    normal = document.styles["Normal"]
    assert normal.font.name == "Times New Roman"
    assert normal.font.size == Pt(12)
    assert normal.paragraph_format.line_spacing == 1.15
    section = document.sections[0]
    assert abs(section.left_margin - Mm(30)) < 2_000
    assert abs(section.right_margin - Mm(18)) < 2_000
    assert section.page_width > section.page_height  # Landscape summary appendix.
    text = "\n".join(
        [
            *(paragraph.text for paragraph in document.paragraphs),
            *(
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ),
        ]
    )
    assert "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in text
    assert "Độc lập - Tự do - Hạnh phúc" in text
    assert "BÁO CÁO TỔNG HỢP SỐ LIỆU VĂN HÓA – XÃ HỘI" in text
    assert all(len(table.columns) <= 6 for table in document.tables)
    with ZipFile(BytesIO(generate_docx_file(period, reports, villages))) as archive:
        footer_xml = b"".join(
            archive.read(name)
            for name in archive.namelist()
            if name.startswith("word/footer")
        )
    assert b'w:instr="PAGE"' in footer_xml

    pdf_bytes = generate_pdf_file(period, reports, villages)
    reader = PdfReader(BytesIO(pdf_bytes))
    pdf_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert len(reader.pages) >= 3
    assert "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in pdf_text
    assert "Độc lập - Tự do - Hạnh phúc" in pdf_text
    assert "CT14" in pdf_text
    assert "Trang 1" in pdf_text


def test_internal_html_preview_escapes_untrusted_cells_and_shows_all_indicators() -> None:
    reports, villages = _report()
    preview = generate_preview_html("<b>period</b>", reports, villages)

    assert "&lt;script&gt;Village&lt;/script&gt;" in preview
    assert "&lt;b&gt;period&lt;/b&gt;" in preview
    assert preview.count("<th>CT") == 14
    assert "<td style=\"padding: 12px 16px; text-align: center;\"></td>" in preview


def test_summary_export_keeps_submitted_report_that_needs_revision() -> None:
    reports = [
        {
            "village_id": "village-1",
            "workflow_status": "needs_revision",
            "timeliness_status": "on_time",
            "submitted_at": "2026-07-15T10:00:00Z",
            "values": {
                code: 320 if code == "CT01" else 1
                for code in (f"CT{i:02d}" for i in range(1, 15))
            },
        }
    ]

    workbook = openpyxl.load_workbook(
        BytesIO(
            generate_summary_xlsx_file(
                "Tháng 7/2026", reports, {"village-1": "Thôn An Sơn"}
            )
        )
    )
    summary = workbook["Bảng tổng hợp"]
    progress = workbook["Theo dõi tiến độ"]

    assert summary["C5"].value == 320
    assert summary["Q5"].value == "Cần chỉnh sửa · Đúng hạn"
    assert summary["C6"].value == 320
    assert summary["Q6"].value == "1/1"
    assert progress["E5"].value == "15/07/2026 10:00"
    assert progress["F5"].value == "Cần chỉnh sửa · Đúng hạn"


def test_semantic_kpis_match_across_exports_without_absorbing_revision_rows() -> None:
    period_id = "period-2026-07"
    reports = [
        {
            "id": "report-a",
            "village_id": "village-a",
            "period_id": period_id,
            "workflow_status": "approved",
            "timeliness_status": "on_time",
            "report_source": "manual",
            "version": 4,
            "submitted_at": "2026-07-10T10:00:00Z",
            "values": {
                "CT01": 100,
                "CT02": 100,
                "CT03": 10,
                "CT04": 5,
                "CT09": 80,
                "CT11": 100,
            },
        },
        {
            "id": "report-b",
            "village_id": "village-b",
            "period_id": period_id,
            "workflow_status": "locked",
            "timeliness_status": "late",
            "report_source": "excel",
            "version": 7,
            "submitted_at": "2026-07-11T10:00:00Z",
            "values": {
                "CT01": 900,
                "CT02": 900,
                "CT03": 90,
                "CT04": 45,
                "CT09": 720,
                "CT11": 450,
            },
        },
        {
            "id": "report-revision",
            "village_id": "village-c",
            "period_id": period_id,
            "workflow_status": "needs_revision",
            "timeliness_status": "on_time",
            "report_source": "manual",
            "version": 2,
            "submitted_at": "2026-07-12T10:00:00Z",
            "values": {
                "CT01": 100_000,
                "CT02": 100_000,
                "CT03": 100_000,
                "CT04": 100_000,
                "CT09": 100_000,
                "CT11": 100_000,
            },
        },
    ]
    villages = {
        "village-a": "Thôn A",
        "village-b": "Thôn B",
        "village-c": "Thôn C",
    }

    workbook = openpyxl.load_workbook(
        BytesIO(
            generate_summary_xlsx_file(
                "Tháng 7/2026",
                reports,
                villages,
                period_id=period_id,
            )
        )
    )
    dashboard = workbook["Dashboard"]
    bhyt_row = next(
        row
        for row in range(1, dashboard.max_row + 1)
        if dashboard.cell(row=row, column=1).value
        == "Tỷ lệ tham gia BHYT"
    )
    assert dashboard.cell(row=bhyt_row, column=2).value == 55
    assert dashboard.cell(row=bhyt_row, column=4).value.startswith(
        "2/3 thôn"
    )
    assert "report-a@v4" in dashboard.cell(
        row=bhyt_row, column=5
    ).value
    assert "report-b@v7" in dashboard.cell(
        row=bhyt_row, column=5
    ).value
    assert "report-revision" not in dashboard.cell(
        row=bhyt_row, column=5
    ).value

    source_sheet = workbook["Bảng tổng hợp"]
    revision_row = next(
        row
        for row in range(5, source_sheet.max_row)
        if source_sheet.cell(row=row, column=2).value == "Thôn C"
    )
    assert source_sheet.cell(row=revision_row, column=3).value == 100_000
    assert "Cần chỉnh sửa" in source_sheet.cell(
        row=revision_row, column=17
    ).value

    document = docx.Document(
        BytesIO(
            generate_docx_file(
                "Tháng 7/2026",
                reports,
                villages,
                period_id=period_id,
            )
        )
    )
    docx_text = "\n".join(
        [
            *(paragraph.text for paragraph in document.paragraphs),
            *(
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ),
        ]
    )
    assert "CHỈ SỐ NGỮ NGHĨA — HỒ SƠ ĐÃ DUYỆT/KHÓA" in docx_text
    assert "55,0%" in docx_text
    assert "2/3 thôn" in docx_text

    pdf_bytes = generate_pdf_file(
        "Tháng 7/2026",
        reports,
        villages,
        period_id=period_id,
    )
    pdf_text = "\n".join(
        page.extract_text() or ""
        for page in PdfReader(BytesIO(pdf_bytes)).pages
    )
    assert "55,0%" in pdf_text
    assert "2/3 thôn" in pdf_text


def test_summary_export_uses_official_ten_village_order() -> None:
    villages = {
        "an-son": "Thôn An Sơn",
        "hoa-ninh": "Thôn Hòa Ninh",
        "phu-hoa": "Thôn Phú Hòa",
        "thach-nham-dong": "Thôn Thạch Nham Đông",
    }

    workbook = openpyxl.load_workbook(
        BytesIO(generate_summary_xlsx_file("Tháng 7/2026", [], villages))
    )
    summary = workbook["Bảng tổng hợp"]

    assert [summary.cell(row=row, column=2).value for row in range(5, 9)] == [
        "Thôn Thạch Nham Đông",
        "Thôn Phú Hòa",
        "Thôn Hòa Ninh",
        "Thôn An Sơn",
    ]


def test_summary_export_indexes_reports_once_and_keeps_first_duplicate() -> None:
    class CountingReport(dict):
        village_id_reads = 0

        def __getitem__(self, key: object) -> object:
            if key == "village_id":
                type(self).village_id_reads += 1
            return super().__getitem__(key)

    report_count = 40
    reports = [
        CountingReport(
            village_id=f"village-{index}",
            values={"CT01": index},
        )
        for index in range(report_count)
    ]
    reports.append(
        CountingReport(
            village_id="village-0",
            values={"CT01": 999},
        )
    )
    villages = {
        f"village-{index}": f"Thôn hiệu năng {index:02d}"
        for index in range(report_count)
    }

    workbook = openpyxl.load_workbook(
        BytesIO(generate_summary_xlsx_file("Tháng 7/2026", reports, villages))
    )
    summary = workbook["Bảng tổng hợp"]

    assert CountingReport.village_id_reads == len(reports)
    assert summary["C5"].value == 0
