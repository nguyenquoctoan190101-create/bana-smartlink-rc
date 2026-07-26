from __future__ import annotations

import html
import hashlib
import io
import json
import logging
import re
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Annotated, Literal
from urllib.parse import quote
from uuid import UUID

import docx
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from fastapi import (
    APIRouter,
    Depends,
    File,
    Form,
    Header,
    HTTPException,
    Request,
    Response,
    UploadFile,
    status,
)
from fastapi.responses import HTMLResponse
from pydantic import BaseModel, Field, field_validator
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from routers.auth import (
    _extract_bearer_token,
    get_settings,
    get_supabase_admin,
    require_admin_or_leader,
    require_admin_xa,
    require_authenticated_user,
)
from services.excel_report_parser import (
    ExcelReportParseError,
    parse_official_report_excel,
)
from services.extraction_review import (
    ExtractionReviewTokenError,
    extraction_values_digest,
    issue_extraction_review_token,
    verify_extraction_review_token,
)
from services.export_service import (
    generate_summary_xlsx_file,
    generate_village_xlsx_file,
)
from services.form_normalizer import (
    FormNormalizationError,
    normalize_excel,
    normalize_field_name,
)
from services.gemini import GeminiError, get_gemini_client
from services.ocr_report import (
    OcrError,
    OcrInputError,
    ocr_report_document_async,
)
from services.operations import RULE_VERSION as VALIDATION_RULE_VERSION
from services.rate_limit import limiter
from services.report_repository import ReportRepository, VillageSubmissionStatus
from services.settings import Settings
from services.supabase_admin import SupabaseAdminClient, SupabaseAdminError, UserProfile
from services.upload_validator import UploadValidationError, validate_report_upload
from services.validator import (
    BLOCKING_ERROR_TYPES,
    ValidationError,
    coerce_storage_value,
    validate_phone,
    validate_report,
)


logger = logging.getLogger(__name__)
router = APIRouter(prefix="/reports", tags=["reports"])
period_router = APIRouter(prefix="/report-periods", tags=["report-periods"])
RULES_PATH = Path(__file__).resolve().parents[1] / "config" / "validation_rules.json"
REPORT_TEMPLATE_VERSION = "ct14-official-2026-07"
REPORT_SOURCE_MAP = {
    "web_form": "manual",
    "manual": "manual",
    "excel_upload": "excel",
    "excel": "excel",
    "photo_upload": "photo_ocr",
    "photo_ocr": "photo_ocr",
    "direct_api": "direct_api",
}
REPORT_PERIOD_CALENDAR_RE = re.compile(
    r"^(?:th[aá]ng\s*)?(\d{1,2})\s*/\s*(\d{4})$",
    re.IGNORECASE,
)


def normalize_report_period_name(value: str) -> str:
    """Normalize harmless whitespace without changing a period's meaning."""

    return " ".join(value.split())


def report_period_name_issue(value: str) -> str | None:
    """Reject impossible months while still allowing descriptive period names."""

    normalized = normalize_report_period_name(value)
    if not normalized:
        return "Tên kỳ báo cáo không được để trống."
    match = REPORT_PERIOD_CALENDAR_RE.fullmatch(normalized)
    if match and not 1 <= int(match.group(1)) <= 12:
        return "Tháng của kỳ báo cáo phải từ 1 đến 12."
    return None


class ExtractionCorrection(BaseModel):
    code: str = Field(pattern=r"^CT(0[1-9]|1[0-4])$")
    before: int | None = Field(default=None, ge=0)
    after: int = Field(ge=0)
    reason: str = Field(min_length=3, max_length=240)

    @field_validator("reason")
    @classmethod
    def normalize_reason(cls, value: str) -> str:
        normalized = " ".join(value.split())
        if len(normalized) < 3:
            raise ValueError("Correction reason is required")
        return normalized


class ExtractionEvidenceReference(BaseModel):
    confidence: float = Field(ge=0, le=1)
    source_page: int | None = Field(default=None, ge=1)
    source_region: str | None = Field(default=None, max_length=240)
    extractor: str = Field(min_length=1, max_length=120)
    method: str = Field(min_length=1, max_length=120)
    version: str = Field(min_length=1, max_length=120)
    flags: list[str] = Field(default_factory=list, max_length=30)
    requires_review: bool


class ExtractionQualitySummary(BaseModel):
    status: Literal["ready", "needs_review", "blocked"]
    mean_confidence: float = Field(ge=0, le=1)
    blocking_flag_count: int = Field(ge=0, le=100)
    warning_flag_count: int = Field(ge=0, le=100)


class ExtractionMetadata(BaseModel):
    source_checksum: str = Field(pattern=r"^[0-9a-f]{64}$")
    source_type: Literal["excel", "photo_ocr", "pdf_ocr"]
    extractor_versions: list[str] = Field(default_factory=list, max_length=20)
    field_count: int = Field(default=14, ge=0, le=14)
    requires_review_count: int = Field(default=0, ge=0, le=14)
    template_version: str | None = Field(default=None, max_length=120)
    rule_version: str | None = Field(default=None, max_length=120)
    evidence_sha256: str | None = Field(default=None, pattern=r"^[0-9a-f]{64}$")
    evidence: dict[str, ExtractionEvidenceReference] = Field(default_factory=dict)
    quality_summary: ExtractionQualitySummary | None = None

    @field_validator("extractor_versions")
    @classmethod
    def normalize_extractor_versions(cls, values: list[str]) -> list[str]:
        normalized = list(
            dict.fromkeys(" ".join(value.split())[:120] for value in values if value.strip())
        )
        if len(normalized) > 20:
            raise ValueError("Too many extractor versions")
        return normalized


class OfflineReportItem(BaseModel):
    id: UUID
    village_id: UUID
    # Offline storage deliberately strips profile PII. The server restores
    # these values from the authenticated profile during synchronization.
    reporter_name: str = Field(default="", max_length=120)
    reporter_phone: str = Field(default="", max_length=20)
    period_id: UUID | None = None
    report_period: str | None = Field(default=None, min_length=1, max_length=120)
    # Kept only for backwards compatibility with older offline queues.  The
    # canonical workflow/timeliness/publication fields are computed server-side.
    status: str | None = None
    updated_at: str
    CT01: int | None = None
    CT02: int | None = None
    CT03: int | None = None
    CT04: int | None = None
    CT05: int | None = None
    CT06: int | None = None
    CT07: int | None = None
    CT08: int | None = None
    CT09: int | None = None
    CT10: int | None = None
    CT11: int | None = None
    CT12: int | None = None
    CT13: int | None = None
    CT14: int | None = None
    assisted_by_cnscd: bool = False
    assisted_member_name: str | None = None
    raw_source: str = "web_form"
    source_confirmed: bool = False
    extraction_corrections: list["ExtractionCorrection"] = Field(default_factory=list, max_length=14)
    extraction_metadata: "ExtractionMetadata | None" = None
    extraction_review_token: str | None = Field(default=None, max_length=8192)
    expected_version: int | None = Field(default=None, ge=1)
    idempotency_key: UUID | None = None


class SyncReportsRequest(BaseModel):
    reports: list[OfflineReportItem]


class AcceptedReportItem(BaseModel):
    client_id: UUID
    report_id: UUID
    version: int
    workflow_status: str
    timeliness_status: str
    publication_status: str


class RejectedReportItem(BaseModel):
    client_id: UUID
    code: str
    message: str
    retryable: bool


class SyncReportsResponse(BaseModel):
    accepted: list[AcceptedReportItem]
    rejected: list[RejectedReportItem]


class ReportSubmitRequest(BaseModel):
    village_id: UUID
    period_id: UUID
    submitted_by_name: str = Field(min_length=1, max_length=120)
    submitted_by_phone: str = Field(min_length=10, max_length=20)
    assisted_by_cnscd: bool = False
    assisted_member_name: str | None = Field(default=None, max_length=120)
    values: dict[str, Any]
    raw_source: str = "web_form"
    source_confirmed: bool = False
    extraction_corrections: list["ExtractionCorrection"] = Field(default_factory=list, max_length=14)
    extraction_metadata: "ExtractionMetadata | None" = None
    extraction_review_token: str | None = Field(default=None, max_length=8192)
    expected_version: int | None = Field(default=None, ge=1)
    idempotency_key: UUID


class ValidationErrorResponse(BaseModel):
    ct_code: str
    error_type: str
    message: str


class ReportSubmitResponse(BaseModel):
    report_id: UUID
    village_id: UUID
    period_id: UUID
    status: str
    workflow_status: str
    timeliness_status: str
    version: int
    replayed: bool = False
    validation_flags: list[ValidationErrorResponse] = Field(default_factory=list)


class ReportUploadResponse(ReportSubmitResponse):
    filename: str
    size_bytes: int


class VillageStatusResponse(BaseModel):
    village_id: UUID
    village_name: str
    old_village_names: list[str]
    report_id: UUID | None
    submitted_at: str | None
    due_date: str | None
    days_late: int
    status: str
    dashboard_color: str


class ReportsStatusResponse(BaseModel):
    period_id: UUID
    villages: list[VillageStatusResponse]


class CreateReportPeriodRequest(BaseModel):
    name: str = Field(min_length=1, max_length=120)
    due_date: datetime
    village_ids: list[UUID] = Field(min_length=1, max_length=100)
    template_name: str | None = Field(default=None, max_length=255)

    @field_validator("name")
    @classmethod
    def validate_name(cls, value: str) -> str:
        normalized = normalize_report_period_name(value)
        issue = report_period_name_issue(normalized)
        if issue:
            raise ValueError(issue)
        return normalized


class ReportPeriodTemplateResponse(BaseModel):
    period_id: UUID
    template_name: str
    template_path: str
    template_sha256: str
    template_size_bytes: int


class OcrValidationFlag(BaseModel):
    ct_code: str
    error_type: str
    message: str


class ReportFieldEvidence(BaseModel):
    raw_value: int | float | str | None = None
    normalized_value: int | None = None
    confidence: float = Field(ge=0, le=1)
    source_page: int | None = Field(default=None, ge=1)
    source_region: str | None = None
    extractor: str
    method: str
    version: str
    flags: list[str] = Field(default_factory=list)
    requires_review: bool


class ReportPreviewMetadata(BaseModel):
    period_name: str | None = None
    village_name: str | None = None
    reporter_name: str | None = None
    reporter_title: str | None = None
    reporter_phone: str | None = None
    deadline: str | None = None


class OcrPreviewResponse(BaseModel):
    """Returned by /reports/ocr-preview for human review BEFORE saving.

    The staff member (can bo) MUST review and confirm these values.
    This endpoint never persists data automatically.
    """

    values: dict[str, int | None]
    raw_values: dict[str, int | float | str | None] = Field(default_factory=dict)
    flags: list[OcrValidationFlag]
    null_codes: list[str]
    filename: str
    size_bytes: int
    source: Literal["excel", "photo_ocr", "pdf_ocr"]
    checksum_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    extractor_versions: list[str] = Field(default_factory=list)
    extraction_review_token: str = Field(min_length=32, max_length=8192)
    evidence: dict[str, ReportFieldEvidence] = Field(default_factory=dict)
    import_metadata: ExtractionMetadata | None = None
    metadata: ReportPreviewMetadata | None = None
    # raw_gemini_text is intentionally excluded from the response model
    # to prevent the AI-generated text from reaching the frontend.


class ReportImportCapabilities(BaseModel):
    excel_preview_enabled: bool = True
    ocr_preview_enabled: bool
    accepted_ocr_types: list[str] = Field(default_factory=list)


def _build_import_metadata(
    *,
    source_checksum: str,
    source_type: Literal["excel", "photo_ocr", "pdf_ocr"],
    extractor_versions: list[str],
    evidence: dict[str, ReportFieldEvidence],
    validation_flags: list[dict[str, Any]],
) -> ExtractionMetadata:
    evidence_references = {
        code: ExtractionEvidenceReference(
            confidence=item.confidence,
            source_page=item.source_page,
            source_region=item.source_region,
            extractor=item.extractor,
            method=item.method,
            version=item.version,
            flags=item.flags,
            requires_review=item.requires_review,
        )
        for code, item in evidence.items()
    }
    evidence_payload = {
        code: item.model_dump(mode="json")
        for code, item in sorted(evidence_references.items())
    }
    evidence_sha256 = hashlib.sha256(
        json.dumps(
            evidence_payload,
            ensure_ascii=False,
            separators=(",", ":"),
            sort_keys=True,
        ).encode("utf-8")
    ).hexdigest()
    blocking_flag_count = sum(
        str(flag.get("error_type")) in BLOCKING_ERROR_TYPES
        for flag in validation_flags
    )
    warning_flag_count = max(0, len(validation_flags) - blocking_flag_count)
    requires_review_count = sum(
        item.requires_review for item in evidence_references.values()
    )
    mean_confidence = (
        round(
            sum(item.confidence for item in evidence_references.values())
            / len(evidence_references),
            4,
        )
        if evidence_references
        else 0.0
    )
    quality_status: Literal["ready", "needs_review", "blocked"] = (
        "blocked"
        if blocking_flag_count
        else "needs_review"
        if requires_review_count or warning_flag_count
        else "ready"
    )
    return ExtractionMetadata(
        source_checksum=source_checksum,
        source_type=source_type,
        extractor_versions=extractor_versions,
        field_count=len(evidence_references),
        requires_review_count=requires_review_count,
        template_version=REPORT_TEMPLATE_VERSION,
        rule_version=VALIDATION_RULE_VERSION,
        evidence_sha256=evidence_sha256,
        evidence=evidence_references,
        quality_summary=ExtractionQualitySummary(
            status=quality_status,
            mean_confidence=mean_confidence,
            blocking_flag_count=blocking_flag_count,
            warning_flag_count=warning_flag_count,
        ),
    )


class ReportNarrativeRequest(BaseModel):
    """Aggregate-only input for optional, non-authoritative AI narration."""

    values: dict[str, Any]
    period_name: str | None = Field(default=None, max_length=120)


class ReportNarrativeResponse(BaseModel):
    """A read-only AI explanation; deterministic validation remains authoritative."""

    is_valid: bool
    errors: list[str] = Field(default_factory=list)
    warnings: list[str] = Field(default_factory=list)
    recommendations: list[str] = Field(default_factory=list)
    source: Literal["gemini", "deterministic"]
    period_name: str | None = None


def get_report_repository(
    supabase: Annotated[SupabaseAdminClient, Depends(get_supabase_admin)],
    authorization: Annotated[str | None, Header()] = None,
) -> ReportRepository:
    if not authorization:
        raise HTTPException(status_code=401, detail="Missing Authorization header")
    scheme, _, token = authorization.partition(" ")
    if scheme.lower() != "bearer" or not token:
        raise HTTPException(status_code=401, detail="Expected Bearer token")
    return ReportRepository(
        supabase.as_user(token),
        admin_supabase=supabase,
    )


@period_router.get("")
async def list_report_periods(
    repository: Annotated[ReportRepository, Depends(get_report_repository)],
    _: Annotated[UserProfile, Depends(require_authenticated_user)],
) -> list[dict[str, Any]]:
    try:
        rows = await repository._supabase._rest_request(
            "GET",
            (
                "/rest/v1/report_periods"
                "?select=id,name,due_date,template_name,template_path,"
                "template_sha256,template_size_bytes,created_at,"
                "report_period_villages(village_id)"
                "&order=due_date.desc,created_at.desc"
            ),
        )
        results: list[dict[str, Any]] = []
        for row in rows:
            assignments = row.get("report_period_villages") or []
            result = {
                key: value
                for key, value in row.items()
                if key != "report_period_villages"
            }
            result["village_ids"] = [
                str(item["village_id"])
                for item in assignments
                if isinstance(item, dict) and item.get("village_id") is not None
            ]
            results.append(result)
        return results
    except SupabaseAdminError as exc:
        raise HTTPException(
            status_code=502, detail="Unable to load report periods"
        ) from exc


@period_router.post("", status_code=status.HTTP_201_CREATED)
async def create_report_period(
    payload: CreateReportPeriodRequest,
    repository: Annotated[ReportRepository, Depends(get_report_repository)],
    _: Annotated[UserProfile, Depends(require_admin_xa)],
) -> dict[str, Any]:
    try:
        rows = await repository._supabase._rest_request(
            "POST",
            "/rest/v1/rpc/create_report_period",
            {
                "p_name": payload.name,
                "p_due_date": payload.due_date.isoformat(),
                "p_village_ids": [str(item) for item in payload.village_ids],
                "p_template_name": payload.template_name.strip()
                if payload.template_name
                else None,
            },
        )
    except SupabaseAdminError as exc:
        if exc.error_code == "42501" or exc.status_code == 403:
            raise HTTPException(
                status_code=403, detail="Only admin can create periods"
            ) from exc
        if exc.error_code == "23505" or exc.status_code == 409:
            raise HTTPException(
                status_code=409, detail="Report period already exists"
            ) from exc
        if exc.error_code == "22023" or exc.status_code == 422:
            raise HTTPException(
                status_code=422, detail="Invalid report period"
            ) from exc
        raise HTTPException(
            status_code=502, detail="Unable to create report period"
        ) from exc
    if not rows:
        raise HTTPException(
            status_code=502, detail="Report period creation returned no result"
        )
    return rows[0]


@period_router.post(
    "/{period_id}/template",
    response_model=ReportPeriodTemplateResponse,
    status_code=status.HTTP_201_CREATED,
)
async def upload_report_period_template(
    period_id: UUID,
    file: Annotated[UploadFile, File(...)],
    repository: Annotated[ReportRepository, Depends(get_report_repository)],
    _: Annotated[UserProfile, Depends(require_admin_xa)],
) -> ReportPeriodTemplateResponse:
    """Validate and store the actual XLSX template in a private bucket."""
    if Path(file.filename or "").suffix.lower() != ".xlsx":
        raise HTTPException(status_code=422, detail="Template must be an XLSX workbook")
    try:
        content = await validate_report_upload(file)
    except UploadValidationError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc

    try:
        period_rows = await repository._supabase._rest_request(
            "GET",
            (f"/rest/v1/report_periods?id=eq.{period_id}&select=id,commune_id&limit=1"),
        )
    except SupabaseAdminError as exc:
        raise HTTPException(
            status_code=502, detail="Unable to verify report period"
        ) from exc
    if not period_rows:
        raise HTTPException(status_code=404, detail="Report period not found")

    digest = hashlib.sha256(content).hexdigest()
    commune_id = str(period_rows[0]["commune_id"])
    object_path = f"{commune_id}/{period_id}/{digest}.xlsx"
    template_name = Path(file.filename or "template.xlsx").name
    try:
        await repository._supabase.upload_storage_object(
            "report-templates",
            object_path,
            content,
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    except SupabaseAdminError as exc:
        # Re-uploading the exact immutable object is idempotent. The database
        # metadata still has to be written after a previous partial request.
        if exc.status_code != 409:
            raise HTTPException(
                status_code=502, detail="Unable to store report template"
            ) from exc

    try:
        updated_rows = await repository._supabase._rest_request(
            "PATCH",
            f"/rest/v1/report_periods?id=eq.{period_id}",
            {
                "template_name": template_name,
                "template_path": object_path,
                "template_sha256": digest,
                "template_size_bytes": len(content),
            },
            prefer="return=representation",
        )
    except SupabaseAdminError as exc:
        raise HTTPException(
            status_code=502, detail="Template stored but metadata update failed"
        ) from exc
    if not updated_rows:
        raise HTTPException(
            status_code=409, detail="Report period changed during template upload"
        )

    return ReportPeriodTemplateResponse(
        period_id=period_id,
        template_name=template_name,
        template_path=object_path,
        template_sha256=digest,
        template_size_bytes=len(content),
    )


@router.get("")
async def list_reports(
    repository: Annotated[ReportRepository, Depends(get_report_repository)],
    current_user: Annotated[UserProfile, Depends(require_authenticated_user)],
    period_id: UUID | None = None,
    village_id: UUID | None = None,
) -> list[dict[str, Any]]:
    """List reports visible to the JWT caller without duplicated profile PII."""
    if village_id is not None:
        await _authorize_village_read(repository, current_user, village_id)
    target_village = village_id
    if current_user.role == "can_bo_thon":
        if not current_user.village_id:
            raise HTTPException(
                status_code=403, detail="User has no village assignment"
            )
        target_village = UUID(current_user.village_id)
    elif current_user.role == "to_cnscd" and village_id is None:
        # The caller-scoped Supabase client keeps this unfiltered query inside
        # the RLS boundary: CNSCĐ can see their own village and every village
        # explicitly recorded in user_village_assignments.
        target_village = None

    query = (
        "/rest/v1/reports"
        "?select=id,village_id,period_id,workflow_status,timeliness_status,"
        "publication_status,report_source,version,submitted_at,"
        "report_values(ct_code,value,note)"
        "&order=submitted_at.desc.nullslast,created_at.desc"
    )
    if period_id is not None:
        query += f"&period_id=eq.{period_id}"
    if target_village is not None:
        query += f"&village_id=eq.{target_village}"
    try:
        rows = await repository._supabase._rest_request("GET", query)
    except SupabaseAdminError as exc:
        raise HTTPException(status_code=502, detail="Unable to load reports") from exc

    result: list[dict[str, Any]] = []
    known_codes = set(_indicator_codes())
    for row in rows:
        item = {key: value for key, value in row.items() if key != "report_values"}
        item["values"] = {
            str(value["ct_code"]): value.get("value")
            for value in row.get("report_values", [])
            if isinstance(value, dict) and value.get("ct_code") in known_codes
        }
        result.append(item)
    return result


@router.post(
    "", response_model=ReportSubmitResponse, status_code=status.HTTP_201_CREATED
)
@limiter.limit("10/minute")
async def submit_report(
    request: Request,
    payload: ReportSubmitRequest,
    repository: Annotated[ReportRepository, Depends(get_report_repository)],
    current_user: Annotated[UserProfile, Depends(require_authenticated_user)],
) -> ReportSubmitResponse:
    """Submit a village report from JSON after validation."""
    await _authorize_report_write(repository, current_user, payload.village_id)
    assisted_by_cnscd, assisted_member_name = _resolve_cnscd_assistance(
        current_user,
        payload.assisted_by_cnscd,
    )
    source = _canonical_report_source(payload.raw_source)
    if source in ("photo_ocr", "excel") and not payload.source_confirmed:
        raise HTTPException(
            status_code=422,
            detail="Dữ liệu từ OCR hoặc Excel bắt buộc phải có xác nhận thủ công (source_confirmed=true)",
        )
    (
        extraction_corrections,
        extraction_metadata,
        extraction_evidence,
    ) = _validate_extraction_review(
        source=source,
        values=payload.values,
        corrections=payload.extraction_corrections,
        metadata=payload.extraction_metadata,
        review_token=payload.extraction_review_token,
        current_user_id=current_user.id,
    )

    return await _submit_report_values(
        repository=repository,
        village_id=payload.village_id,
        period_id=payload.period_id,
        submitted_by_name=current_user.display_name or payload.submitted_by_name,
        submitted_by_phone=current_user.phone or payload.submitted_by_phone,
        values=payload.values,
        notes=None,
        raw_source=source,
        assisted_by_cnscd=assisted_by_cnscd,
        assisted_member_name=assisted_member_name,
        expected_version=payload.expected_version,
        idempotency_key=payload.idempotency_key,
        extraction_corrections=extraction_corrections,
        extraction_metadata=extraction_metadata,
        extraction_evidence=extraction_evidence,
    )


@router.post("/sync", response_model=SyncReportsResponse)
@limiter.limit("5/minute")
async def sync_reports(
    request: Request,
    payload: SyncReportsRequest,
    repository: Annotated[ReportRepository, Depends(get_report_repository)],
    current_user: Annotated[UserProfile, Depends(require_authenticated_user)],
) -> SyncReportsResponse:
    """Sync multiple reports without deleting client items that were rejected."""
    accepted: list[AcceptedReportItem] = []
    rejected: list[RejectedReportItem] = []

    for report in payload.reports:
        try:
            await _authorize_report_write(repository, current_user, report.village_id)
        except HTTPException as exc:
            rejected.append(
                RejectedReportItem(
                    client_id=report.id,
                    code="FORBIDDEN",
                    message=str(exc.detail),
                    retryable=False,
                )
            )
            continue

        period_id = str(report.period_id) if report.period_id else None
        if period_id is None and report.report_period:
            try:
                period_id = await repository.get_period_id_by_name(report.report_period)
            except SupabaseAdminError:
                rejected.append(
                    RejectedReportItem(
                        client_id=report.id,
                        code="UPSTREAM_UNAVAILABLE",
                        message="Unable to resolve report period",
                        retryable=True,
                    )
                )
                continue
        if period_id is None:
            rejected.append(
                RejectedReportItem(
                    client_id=report.id,
                    code="PERIOD_NOT_FOUND",
                    message="Report period does not exist",
                    retryable=False,
                )
            )
            continue

        values = {
            f"CT{i:02d}": getattr(report, f"CT{i:02d}")
            for i in range(1, 15)
            if getattr(report, f"CT{i:02d}") is not None
        }
        try:
            assisted_by_cnscd, assisted_member_name = _resolve_cnscd_assistance(
                current_user,
                report.assisted_by_cnscd,
            )
            source = _canonical_report_source(report.raw_source)
            if source in {"excel", "photo_ocr"} and not report.source_confirmed:
                raise HTTPException(
                    status_code=422,
                    detail="Imported data requires explicit human confirmation",
                )
            (
                extraction_corrections,
                extraction_metadata,
                extraction_evidence,
            ) = _validate_extraction_review(
                source=source,
                values=values,
                corrections=report.extraction_corrections,
                metadata=report.extraction_metadata,
                review_token=report.extraction_review_token,
                current_user_id=current_user.id,
            )
            submitted = await _submit_report_values(
                repository=repository,
                village_id=report.village_id,
                period_id=UUID(period_id),
                submitted_by_name=current_user.display_name or report.reporter_name,
                submitted_by_phone=current_user.phone or report.reporter_phone,
                values=values,
                notes=None,
                raw_source=source,
                assisted_by_cnscd=assisted_by_cnscd,
                assisted_member_name=assisted_member_name,
                report_id=report.id,
                expected_version=report.expected_version,
                idempotency_key=report.idempotency_key or report.id,
                extraction_corrections=extraction_corrections,
                extraction_metadata=extraction_metadata,
                extraction_evidence=extraction_evidence,
            )
            accepted.append(
                AcceptedReportItem(
                    client_id=report.id,
                    report_id=submitted.report_id,
                    version=submitted.version,
                    workflow_status=submitted.workflow_status,
                    timeliness_status=submitted.timeliness_status,
                    publication_status="private",
                )
            )
        except HTTPException as exc:
            rejected.append(_sync_rejection(report.id, exc))
        except Exception:
            logger.exception("Unexpected offline sync failure")
            rejected.append(
                RejectedReportItem(
                    client_id=report.id,
                    code="INTERNAL_ERROR",
                    message="Unable to sync report",
                    retryable=True,
                )
            )

    return SyncReportsResponse(accepted=accepted, rejected=rejected)


@router.delete(
    "/{report_id}",
    status_code=status.HTTP_204_NO_CONTENT,
    response_class=Response,
)
async def delete_report(
    report_id: UUID,
    expected_version: int,
    repository: Annotated[ReportRepository, Depends(get_report_repository)],
    current_user: Annotated[UserProfile, Depends(require_authenticated_user)],
) -> Response:
    """Delete an own-village draft or a mutable private report as admin."""
    try:
        rows = await repository._supabase._rest_request(
            "GET",
            (
                f"/rest/v1/reports?id=eq.{report_id}"
                "&select=id,village_id,workflow_status,publication_status,version"
            ),
        )
        if not rows:
            raise HTTPException(status_code=404, detail="Report not found")
        report = rows[0]
        if current_user.role != "admin_xa":
            await _authorize_report_write(
                repository,
                current_user,
                UUID(str(report["village_id"])),
            )
        if int(report["version"]) != expected_version:
            raise HTTPException(status_code=409, detail="Report version conflict")
        if (
            report["workflow_status"] == "locked"
            or report["publication_status"] == "published"
        ):
            raise HTTPException(
                status_code=409,
                detail="Locked or published reports cannot be deleted",
            )
        if current_user.role != "admin_xa" and report["workflow_status"] != "draft":
            raise HTTPException(
                status_code=409, detail="Only draft reports can be deleted"
            )

        deleted = await repository._supabase._rest_request(
            "POST",
            "/rest/v1/rpc/delete_report_submission",
            {
                "p_report_id": str(report_id),
                "p_expected_version": expected_version,
            },
        )
        if not deleted:
            raise HTTPException(
                status_code=409, detail="Report changed before deletion"
            )
    except SupabaseAdminError as exc:
        if exc.error_code in {"40001", "42501"} or exc.status_code in {403, 409}:
            raise HTTPException(
                status_code=409,
                detail="Report changed or is not deletable",
            ) from exc
        raise HTTPException(status_code=502, detail="Unable to delete report") from exc
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@router.post("/normalize")
@limiter.limit("15/minute")
async def normalize_report_excel(
    request: Request,
    repository: Annotated[ReportRepository, Depends(get_report_repository)],
    current_user: Annotated[UserProfile, Depends(require_authenticated_user)],
    file: UploadFile = File(...),
) -> dict:
    _ = current_user
    try:
        content = await validate_report_upload(file)
        runtime_synonyms = await repository.field_synonyms()
        normalized = normalize_excel(content, synonyms=runtime_synonyms)
        return {"success": True, "normalized_data": normalized}
    except UploadValidationError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except FormNormalizationError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except SupabaseAdminError as exc:
        raise HTTPException(
            status_code=502,
            detail="Không thể tải ánh xạ chỉ tiêu.",
        ) from exc
    except Exception:
        logger.exception("Excel normalization failed")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Không thể chuẩn hóa tệp Excel.",
        )


@router.post("/confirm-synonym")
@limiter.limit("30/minute")
async def confirm_field_synonym(
    request: Request,
    payload: dict,
    repository: Annotated[ReportRepository, Depends(get_report_repository)],
    _: Annotated[UserProfile, Depends(require_admin_xa)],
) -> dict:
    original_name = payload.get("original_name")
    ct_code = payload.get("ct_code")
    if (
        not isinstance(original_name, str)
        or not original_name.strip()
        or ct_code not in _indicator_codes()
    ):
        raise HTTPException(status_code=400, detail="Thiếu original_name hoặc ct_code")

    try:
        mapping = await repository.confirm_field_synonym(
            original_name.strip(),
            normalize_field_name(original_name),
            ct_code,
        )
    except SupabaseAdminError as exc:
        raise HTTPException(
            status_code=502,
            detail="Không thể lưu ánh xạ chỉ tiêu.",
        ) from exc
    return {"success": True, "mapping": mapping}


@router.post(
    "/upload", response_model=ReportUploadResponse, status_code=status.HTTP_201_CREATED
)
@limiter.limit("20/minute")
async def upload_report_file(
    request: Request,
    repository: Annotated[ReportRepository, Depends(get_report_repository)],
    current_user: Annotated[UserProfile, Depends(require_authenticated_user)],
    village_id: Annotated[UUID, Form()],
    period_id: Annotated[UUID, Form()],
    submitted_by_name: Annotated[str, Form(min_length=1, max_length=120)],
    submitted_by_phone: Annotated[str, Form(min_length=10, max_length=20)],
    idempotency_key: Annotated[UUID, Form()],
    assisted_by_cnscd: Annotated[bool, Form()] = False,
    assisted_member_name: Annotated[str | None, Form(max_length=120)] = None,
    source_confirmed: Annotated[bool, Form()] = False,
    file: UploadFile = File(...),
) -> ReportUploadResponse:
    """Retired mutation path; every import must use preview then canonical submit."""
    _ = (
        request,
        repository,
        current_user,
        village_id,
        period_id,
        submitted_by_name,
        submitted_by_phone,
        idempotency_key,
        assisted_by_cnscd,
        assisted_member_name,
        source_confirmed,
        file,
    )
    raise HTTPException(
        status_code=status.HTTP_410_GONE,
        detail=(
            "Đường tải và lưu trực tiếp đã ngừng hoạt động. "
            "Hãy dùng /reports/excel-preview, rà soát 14 chỉ tiêu, "
            "sau đó gửi qua POST /reports."
        ),
    )


@router.get("/capabilities", response_model=ReportImportCapabilities)
async def report_import_capabilities(
    _: Annotated[UserProfile, Depends(require_authenticated_user)],
    settings: Annotated[Settings, Depends(get_settings)],
) -> ReportImportCapabilities:
    """Return only backend-confirmed import features used to render the UI."""

    ocr_enabled = bool(settings.feature_external_ocr)
    return ReportImportCapabilities(
        ocr_preview_enabled=ocr_enabled,
        accepted_ocr_types=[".jpg", ".jpeg", ".png", ".pdf"] if ocr_enabled else [],
    )


@router.post("/ocr-preview", response_model=OcrPreviewResponse)
@limiter.limit("10/minute")
async def ocr_photo_preview(
    request: Request,
    current_user: Annotated[UserProfile, Depends(require_authenticated_user)],
    settings: Annotated[Settings, Depends(get_settings)],
    file: UploadFile = File(...),
) -> OcrPreviewResponse:
    """Experimental OCR preview; disabled unless explicitly enabled in dev/test."""
    _ = request
    if not settings.feature_external_ocr:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=(
                "Nhận dạng ảnh/PDF đang khóa để bảo vệ dữ liệu cá nhân. "
                "Vui lòng dùng biểu mẫu Excel hoặc nhập trực tiếp."
            ),
        )
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in {".jpg", ".jpeg", ".png", ".pdf"}:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="OCR chỉ nhận ảnh .jpg, .png hoặc PDF quét.",
        )

    try:
        content = await validate_report_upload(file)
    except UploadValidationError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(exc),
        ) from exc

    try:
        preview = await ocr_report_document_async(content)
    except OcrInputError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(exc),
        ) from exc
    except OcrError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail="OCR processing failed",
        ) from exc

    source_type = "pdf_ocr" if suffix == ".pdf" else "photo_ocr"
    source_checksum = hashlib.sha256(content).hexdigest()
    extractor_versions = sorted(
        {
            f"{item.extractor}:{item.version}"
            for item in preview.evidence.values()
        }
    )
    requires_review_count = sum(
        bool(item.requires_review) for item in preview.evidence.values()
    )
    evidence_models = {
        code: ReportFieldEvidence(
            raw_value=item.raw_value,
            normalized_value=item.normalized_value,
            confidence=item.confidence,
            source_page=item.source_page,
            source_region=item.source_region,
            extractor=item.extractor,
            method=item.method,
            version=item.version,
            flags=item.flags,
            requires_review=item.requires_review,
        )
        for code, item in preview.evidence.items()
    }
    import_metadata = _build_import_metadata(
        source_checksum=source_checksum,
        source_type=source_type,
        extractor_versions=extractor_versions,
        evidence=evidence_models,
        validation_flags=list(preview.flags),
    )
    try:
        review_token = issue_extraction_review_token(
            user_id=current_user.id,
            source_checksum=source_checksum,
            source_type=source_type,
            extractor_versions=extractor_versions,
            values=preview.values,
            requires_review_count=requires_review_count,
            import_metadata=import_metadata.model_dump(mode="json"),
        )
    except ExtractionReviewTokenError as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Không thể tạo bằng chứng rà soát an toàn",
        ) from exc

    return OcrPreviewResponse(
        values=preview.values,
        raw_values=preview.raw_values,
        flags=[
            OcrValidationFlag(
                ct_code=f["ct_code"],
                error_type=f["error_type"],
                message=f["message"],
            )
            for f in preview.flags
        ],
        null_codes=preview.null_codes,
        filename=file.filename or "upload.jpg",
        size_bytes=len(content),
        source=source_type,
        checksum_sha256=source_checksum,
        extractor_versions=extractor_versions,
        extraction_review_token=review_token,
        evidence=evidence_models,
        import_metadata=import_metadata,
        # raw_gemini_text deliberately omitted from response
    )


@router.post("/ai-narrative", response_model=ReportNarrativeResponse)
@limiter.limit("5/minute")
async def create_report_narrative(
    request: Request,
    payload: ReportNarrativeRequest,
    _: Annotated[UserProfile, Depends(require_authenticated_user)],
) -> ReportNarrativeResponse:
    """Explain an aggregate report without persisting or deciding its validity.

    Only CT01-CT14 values and deterministic validation flags may leave the
    application.  Names, phones, villages, GPS, report identifiers and any
    other personal/internal context are intentionally excluded from the prompt.
    """
    _ = request
    indicator_codes = {f"CT{index:02d}" for index in range(1, 15)}
    unknown_codes = sorted(set(payload.values) - indicator_codes)
    if unknown_codes:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail="Chỉ chấp nhận các chỉ tiêu CT01 đến CT14 cho diễn giải AI.",
        )

    aggregate_values = {
        code: payload.values.get(code) for code in sorted(indicator_codes)
    }
    flags = validate_report(aggregate_values)
    blocking_messages = [
        flag["message"] for flag in flags if flag["error_type"] in BLOCKING_ERROR_TYPES
    ]
    warning_messages = [
        flag["message"]
        for flag in flags
        if flag["error_type"] not in BLOCKING_ERROR_TYPES
    ]
    if blocking_messages:
        return ReportNarrativeResponse(
            is_valid=False,
            errors=blocking_messages,
            warnings=warning_messages,
            recommendations=[
                "Hoàn thiện các chỉ tiêu được nêu trước khi yêu cầu diễn giải."
            ],
            source="deterministic",
            period_name=payload.period_name,
        )

    response_schema = {
        "type": "OBJECT",
        "properties": {
            "warnings": {"type": "ARRAY", "items": {"type": "STRING"}},
            "recommendations": {"type": "ARRAY", "items": {"type": "STRING"}},
        },
        "required": ["warnings", "recommendations"],
    }
    system_prompt = (
        "Bạn là trợ lý diễn giải dữ liệu cho hệ thống điều hành cấp xã. "
        "Chỉ dùng các số tổng hợp CT01-CT14 được cung cấp. Không suy đoán, "
        "không tự sửa số, không kết luận báo cáo hợp lệ và không đề xuất công bố. "
        "Không nhắc đến dữ liệu cá nhân, CT14 hoặc thông tin nội bộ. "
        "Nêu tối đa 3 cảnh báo và 3 gợi ý kiểm tra ngắn, trung tính, bằng tiếng Việt."
    )
    # CT14 is internal-only. Validate it locally but never send it to an AI
    # provider, including when the caller is an authenticated staff member.
    ai_values = {
        code: aggregate_values[code] for code in sorted(indicator_codes - {"CT14"})
    }
    user_text = json.dumps(
        {
            "period": payload.period_name or "Chưa nêu kỳ",
            "values": ai_values,
            "known_warnings": warning_messages,
        },
        ensure_ascii=False,
    )
    try:
        generated = await get_gemini_client().generate_json(
            system_prompt=system_prompt,
            user_text=user_text,
            response_schema=response_schema,
        )
    except GeminiError as exc:
        logger.warning("AI narrative unavailable: %s", type(exc).__name__)
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Dịch vụ diễn giải AI hiện chưa sẵn sàng. Kiểm tra nghiệp vụ vẫn hoạt động bình thường.",
        ) from exc

    def clean_items(value: Any) -> list[str]:
        if not isinstance(value, list):
            return []
        return [
            item.strip()[:300]
            for item in value
            if isinstance(item, str) and item.strip()
        ][:3]

    return ReportNarrativeResponse(
        is_valid=True,
        warnings=warning_messages + clean_items(generated.get("warnings")),
        recommendations=clean_items(generated.get("recommendations")),
        source="gemini",
        period_name=payload.period_name,
    )


@router.post("/excel-preview", response_model=OcrPreviewResponse)
@limiter.limit("20/minute")
async def excel_preview(
    request: Request,
    current_user: Annotated[UserProfile, Depends(require_authenticated_user)],
    file: UploadFile = File(...),
) -> OcrPreviewResponse:
    """Parse an official Excel template and return a preview for human review.

    This endpoint NEVER saves data. It only parses and returns values so
    the can_bo can review before submitting via POST /reports.
    """
    _ = request
    if Path(file.filename or "").suffix.lower() != ".xlsx":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Excel preview chỉ nhận file .xlsx.",
        )

    try:
        content = await validate_report_upload(file)
    except UploadValidationError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(exc),
        ) from exc

    try:
        parsed = parse_official_report_excel(content)
    except ExcelReportParseError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(exc),
        ) from exc

    raw_values = parsed["values"]
    normalized_values = {
        code: coerce_storage_value(value) for code, value in raw_values.items()
    }
    validation_flags = validate_report(raw_values)
    reporter_phone = parsed["metadata"].get("reporter_phone")
    phone_flag = validate_phone(reporter_phone) if reporter_phone is not None else None
    if phone_flag is not None:
        validation_flags.append(phone_flag)
    null_codes = [code for code, value in normalized_values.items() if value is None]
    flag_types_by_code: dict[str, list[str]] = {
        code: [] for code in normalized_values
    }
    for flag in validation_flags:
        code = flag["ct_code"]
        if code in flag_types_by_code:
            flag_types_by_code[code].append(flag["error_type"])

    source_checksum = hashlib.sha256(content).hexdigest()
    extractor_versions = sorted(
        {
            f"{item['extractor']}:{item['version']}"
            for item in parsed["evidence"].values()
        }
    )
    requires_review_count = sum(
        item["requires_review"] or bool(flag_types_by_code[code])
        for code, item in parsed["evidence"].items()
    )
    evidence_models = {
        code: ReportFieldEvidence(
            raw_value=item["raw_value"],
            normalized_value=item["normalized_value"],
            confidence=item["confidence"],
            source_page=item["source_page"],
            source_region=item["source_region"],
            extractor=item["extractor"],
            method=item["method"],
            version=item["version"],
            flags=list(
                dict.fromkeys(item["flags"] + flag_types_by_code[code])
            ),
            requires_review=(
                item["requires_review"] or bool(flag_types_by_code[code])
            ),
        )
        for code, item in parsed["evidence"].items()
    }
    import_metadata = _build_import_metadata(
        source_checksum=source_checksum,
        source_type="excel",
        extractor_versions=extractor_versions,
        evidence=evidence_models,
        validation_flags=list(validation_flags),
    )
    try:
        review_token = issue_extraction_review_token(
            user_id=current_user.id,
            source_checksum=source_checksum,
            source_type="excel",
            extractor_versions=extractor_versions,
            values=normalized_values,
            requires_review_count=requires_review_count,
            import_metadata=import_metadata.model_dump(mode="json"),
        )
    except ExtractionReviewTokenError as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Không thể tạo bằng chứng rà soát an toàn",
        ) from exc

    return OcrPreviewResponse(
        values=normalized_values,
        raw_values=raw_values,
        flags=[OcrValidationFlag(**flag) for flag in validation_flags],
        null_codes=null_codes,
        filename=file.filename or "upload.xlsx",
        size_bytes=len(content),
        source="excel",
        checksum_sha256=source_checksum,
        extractor_versions=extractor_versions,
        extraction_review_token=review_token,
        evidence=evidence_models,
        import_metadata=import_metadata,
        metadata=ReportPreviewMetadata(**parsed["metadata"]),
    )


@router.get("/periods")
async def get_report_periods(
    supabase: Annotated[SupabaseAdminClient, Depends(get_supabase_admin)],
    _: Annotated[UserProfile, Depends(require_authenticated_user)],
    authorization: Annotated[str | None, Header()] = None,
):
    """Return all report periods."""
    caller = supabase.as_user(_extract_bearer_token(authorization))
    try:
        return await caller._rest_request(
            "GET",
            "/rest/v1/report_periods?select=id,name,due_date&order=due_date.desc",
        )
    except SupabaseAdminError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail="Không lấy được danh sách kỳ báo cáo.",
        ) from exc


@router.get("/villages")
async def get_villages(
    supabase: Annotated[SupabaseAdminClient, Depends(get_supabase_admin)],
    settings: Annotated[Settings, Depends(get_settings)],
):
    """Return the active canonical village catalogue from the database."""
    try:
        commune_id = settings.bana_commune_id
        encoded_commune_id = quote(commune_id, safe="")
        rows = await supabase._rest_request(
            "GET",
            (
                f"/rest/v1/villages?commune_id=eq.{encoded_commune_id}"
                "&is_active=eq.true&select=id,name,commune_id&order=name.asc"
            ),
        )
        return [
            {"id": row["id"], "name": row["name"]}
            for row in rows
            if str(row.get("commune_id", "")) == commune_id
            and row.get("id") is not None
            and isinstance(row.get("name"), str)
        ]
    except SupabaseAdminError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail="Không lấy được danh mục thôn.",
        ) from exc


@router.get("/public")
async def get_public_reports(
    supabase: Annotated[SupabaseAdminClient, Depends(get_supabase_admin)],
    settings: Annotated[Settings, Depends(get_settings)],
):
    """Return public reports for citizens (filtering sensitive CT codes)."""
    try:
        commune_id = settings.bana_commune_id
        encoded_commune_id = quote(commune_id, safe="")
        public_codes = {"CT01", "CT02", "CT09", "CT12", "CT13"}
        codes_filter = ",".join(sorted(public_codes))
        reports = await supabase._rest_request(
            "GET",
            (
                "/rest/v1/reports?publication_status=eq.published"
                "&select=id,village_id,published_at,"
                "report_periods!inner(name,commune_id),"
                "villages!inner(commune_id),"
                "report_values!inner(ct_code,value)"
                f"&report_periods.commune_id=eq.{encoded_commune_id}"
                f"&villages.commune_id=eq.{encoded_commune_id}"
                f"&report_values.ct_code=in.({codes_filter})"
                "&order=published_at.desc"
            ),
        )

        # Format response
        result = []
        for r in reports:
            period_scope = r.get("report_periods")
            village_scope = r.get("villages")
            if (
                not isinstance(period_scope, dict)
                or str(period_scope.get("commune_id", "")) != commune_id
                or not isinstance(village_scope, dict)
                or str(village_scope.get("commune_id", "")) != commune_id
            ):
                continue
            values_dict = {
                v["ct_code"]: v["value"]
                for v in r.get("report_values", [])
                if v["ct_code"] in public_codes
            }

            result.append(
                {
                    "id": r["id"],
                    "village_id": r["village_id"],
                    "report_period": period_scope.get("name", "Unknown"),
                    "published_at": r.get("published_at"),
                    "values": values_dict,
                }
            )
        return result
    except SupabaseAdminError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail="Không lấy được dữ liệu công khai.",
        ) from exc


@router.get("/status", response_model=ReportsStatusResponse)
async def get_reports_status(
    period_id: str,
    repository: Annotated[ReportRepository, Depends(get_report_repository)],
    _: Annotated[UserProfile, Depends(require_authenticated_user)],
) -> ReportsStatusResponse:
    """Return submission status for all current villages in one period."""
    resolved_uuid, _ = await safe_resolve_period(repository._supabase, period_id)
    try:
        statuses = await repository.submission_statuses(str(resolved_uuid))
    except SupabaseAdminError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail="Không lấy được trạng thái nộp báo cáo.",
        ) from exc

    return ReportsStatusResponse(
        period_id=resolved_uuid,
        villages=[_status_response(item) for item in statuses],
    )


class TrendAlertResponse(BaseModel):
    village_id: UUID
    village_name: str
    ct_code: str
    indicator_name: str
    prev_value: int
    curr_value: int
    change_pct: float


@router.get("/trend-alerts", response_model=list[TrendAlertResponse])
async def get_trend_alerts(
    curr_period_id: str,
    prev_period_id: str,
    supabase: Annotated[SupabaseAdminClient, Depends(get_supabase_admin)],
    _: Annotated[UserProfile, Depends(require_admin_or_leader)],
    authorization: Annotated[str | None, Header()] = None,
) -> list[TrendAlertResponse]:
    """Compare reports from two periods and return indicators with >20% changes."""
    from services.trend_alert import get_trend_alerts_async  # noqa: PLC0415

    caller = supabase.as_user(_extract_bearer_token(authorization))
    curr_uuid, _ = await safe_resolve_period(caller, curr_period_id)
    prev_uuid, _ = await safe_resolve_period(caller, prev_period_id)
    try:
        alerts = await get_trend_alerts_async(
            supabase=caller,
            prev_period_id=str(prev_uuid),
            curr_period_id=str(curr_uuid),
        )
    except Exception as exc:
        logger.exception("Trend analysis failed")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Unable to analyze report trends",
        ) from exc

    return [
        TrendAlertResponse(
            village_id=UUID(a["village_id"]),
            village_name=a["village_name"],
            ct_code=a["ct_code"],
            indicator_name=a["indicator_name"],
            prev_value=a["prev_value"],
            curr_value=a["curr_value"],
            change_pct=a["change_pct"],
        )
        for a in alerts
    ]


async def _submit_report_values(
    repository: ReportRepository,
    village_id: UUID,
    period_id: UUID,
    submitted_by_name: str,
    submitted_by_phone: str,
    values: dict[str, Any],
    notes: dict[str, str | None] | None,
    raw_source: str,
    assisted_by_cnscd: bool = False,
    assisted_member_name: str | None = None,
    report_id: UUID | None = None,
    expected_version: int | None = None,
    idempotency_key: UUID | None = None,
    extraction_corrections: list[dict[str, Any]] | None = None,
    extraction_metadata: dict[str, Any] | None = None,
    extraction_evidence: dict[str, Any] | None = None,
) -> ReportSubmitResponse:
    unknown_codes = sorted(set(values) - set(_indicator_codes()))
    if unknown_codes:
        raise HTTPException(
            status_code=422,
            detail={
                "errors": [
                    {
                        "ct_code": code,
                        "error_type": "TEXT",
                        "message": f"Unknown indicator code: {code}",
                    }
                    for code in unknown_codes
                ]
            },
        )
    known_values = _known_report_values(values)
    validation_errors = validate_report(known_values)
    phone_error = validate_phone(submitted_by_phone)
    if phone_error is not None:
        validation_errors.append(phone_error)

    if _has_blocking_errors(validation_errors):
        raise HTTPException(
            status_code=422,
            detail={"errors": validation_errors},
        )

    storage_values = {
        ct_code: coerce_storage_value(value) for ct_code, value in known_values.items()
    }
    non_blocking_flags = [
        error
        for error in validation_errors
        if error["error_type"] not in BLOCKING_ERROR_TYPES
    ]

    try:
        saved_report = await repository.save_report(
            village_id=str(village_id),
            period_id=str(period_id),
            submitted_by_name=submitted_by_name.strip(),
            submitted_by_phone=submitted_by_phone.strip(),
            values=storage_values,
            flags=non_blocking_flags,
            raw_source=raw_source,  # type: ignore[arg-type]
            notes=notes,
            assisted_by_cnscd=assisted_by_cnscd,
            assisted_member_name=assisted_member_name,
            report_id=str(report_id) if report_id else None,
            expected_version=expected_version,
            idempotency_key=str(idempotency_key) if idempotency_key else None,
            extraction_corrections=extraction_corrections,
            extraction_metadata=extraction_metadata,
            extraction_evidence=extraction_evidence,
        )
    except SupabaseAdminError as exc:
        if exc.error_code == "40001" or exc.status_code == 409:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail="Report version conflict",
            ) from exc
        if exc.error_code == "42501" or exc.status_code == 403:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Not allowed to modify this report",
            ) from exc
        if exc.error_code == "22023" or exc.status_code == 422:
            raise HTTPException(
                status_code=422,
                detail="Invalid report submission",
            ) from exc
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail="Không lưu được báo cáo.",
        ) from exc

    return ReportSubmitResponse(
        report_id=UUID(saved_report.id),
        village_id=UUID(saved_report.village_id),
        period_id=UUID(saved_report.period_id),
        status=saved_report.workflow_status,
        workflow_status=saved_report.workflow_status,
        timeliness_status=saved_report.timeliness_status,
        version=saved_report.version,
        replayed=saved_report.replayed,
        validation_flags=non_blocking_flags,
    )


def _validate_extraction_review(
    *,
    source: str,
    values: dict[str, Any],
    corrections: list[ExtractionCorrection],
    metadata: ExtractionMetadata | None,
    review_token: str | None,
    current_user_id: str,
) -> tuple[
    list[dict[str, Any]],
    dict[str, Any] | None,
    dict[str, Any] | None,
]:
    imported = source in {"excel", "photo_ocr"}
    if not imported and (corrections or metadata is not None or review_token is not None):
        raise HTTPException(
            status_code=422,
            detail="Extraction review metadata is only allowed for imported reports",
        )
    if imported and (metadata is None or review_token is None):
        raise HTTPException(
            status_code=422,
            detail="Imported reports require signed extraction review evidence",
        )
    if not imported:
        return [], None, None

    expected_metadata_sources = (
        {"excel"} if source == "excel" else {"photo_ocr", "pdf_ocr"}
    )
    if metadata is None or metadata.source_type not in expected_metadata_sources:
        raise HTTPException(
            status_code=422,
            detail="Extraction source metadata does not match the report source",
        )
    try:
        trusted = verify_extraction_review_token(
            review_token or "",
            user_id=current_user_id,
        )
    except ExtractionReviewTokenError as exc:
        raise HTTPException(
            status_code=422,
            detail="Extraction review evidence is invalid or expired",
        ) from exc
    if trusted["source_type"] not in expected_metadata_sources:
        raise HTTPException(
            status_code=422,
            detail="Signed extraction source does not match the report source",
        )

    core_trusted_metadata = {
        "source_checksum": trusted["source_checksum"],
        "source_type": trusted["source_type"],
        "extractor_versions": trusted["extractor_versions"],
        "field_count": trusted["field_count"],
        "requires_review_count": trusted["requires_review_count"],
    }
    signed_import_metadata = trusted.get("import_metadata")
    if signed_import_metadata is not None:
        if any(
            signed_import_metadata.get(key) != value
            for key, value in core_trusted_metadata.items()
        ):
            raise HTTPException(
                status_code=422,
                detail="Signed import metadata is inconsistent",
            )
        trusted_metadata = signed_import_metadata
    else:
        trusted_metadata = core_trusted_metadata
    supplied_metadata = metadata.model_dump(exclude_none=True)
    if signed_import_metadata is None and not supplied_metadata.get("evidence"):
        supplied_metadata.pop("evidence", None)
    supplied_metadata["extractor_versions"] = sorted(
        supplied_metadata["extractor_versions"]
    )
    if supplied_metadata != trusted_metadata:
        raise HTTPException(
            status_code=422,
            detail="Extraction metadata does not match the signed preview",
        )

    seen: set[str] = set()
    supplied_corrections: dict[str, ExtractionCorrection] = {}
    for correction in corrections:
        if correction.code in seen:
            raise HTTPException(
                status_code=422,
                detail=f"Duplicate extraction correction for {correction.code}",
            )
        seen.add(correction.code)
        supplied_corrections[correction.code] = correction

    normalized: list[dict[str, Any]] = []
    for code, original_value in trusted["values"].items():
        submitted_value = coerce_storage_value(values.get(code))
        correction = supplied_corrections.get(code)
        if submitted_value == original_value:
            if correction is not None:
                raise HTTPException(
                    status_code=422,
                    detail=f"Unchanged extraction field cannot have a correction for {code}",
                )
            continue
        if (
            correction is None
            or correction.before != original_value
            or correction.after != submitted_value
            or correction.before == correction.after
        ):
            raise HTTPException(
                status_code=422,
                detail=f"Every changed extraction field requires before/after/reason for {code}",
            )
        normalized.append(correction.model_dump())

    evidence_registration = {
        "id": trusted["jti"],
        "user_id": current_user_id,
        "source_type": trusted["source_type"],
        "source_checksum": trusted["source_checksum"],
        "extractor_versions": trusted["extractor_versions"],
        "original_values_sha256": extraction_values_digest(trusted["values"]),
        "field_count": trusted["field_count"],
        "requires_review_count": trusted["requires_review_count"],
        "expires_at": datetime.fromtimestamp(trusted["exp"], tz=UTC).isoformat(),
    }
    return normalized, trusted_metadata, evidence_registration


def _canonical_report_source(raw_source: str) -> str:
    source = REPORT_SOURCE_MAP.get(raw_source.strip().lower())
    if source is None:
        raise HTTPException(
            status_code=422,
            detail="Unsupported report source",
        )
    return source


async def _is_assigned_cnscd_village(
    repository: ReportRepository,
    user: UserProfile,
    village_id: UUID,
) -> bool:
    """Check an explicit CNSCĐ assignment using the caller's JWT and RLS."""
    try:
        rows = await repository._supabase._rest_request(
            "GET",
            (
                "/rest/v1/user_village_assignments"
                "?select=village_id"
                f"&user_id=eq.{quote(str(user.id), safe='')}"
                f"&village_id=eq.{village_id}"
                "&limit=1"
            ),
        )
    except SupabaseAdminError as exc:
        raise HTTPException(
            status_code=503,
            detail="Unable to verify village assignment",
        ) from exc
    return bool(rows)


async def _authorize_report_write(
    repository: ReportRepository,
    user: UserProfile,
    village_id: UUID,
) -> None:
    if user.role == "lanh_dao":
        raise HTTPException(status_code=403, detail="Leadership role is read-only")
    if user.role == "admin_xa":
        raise HTTPException(
            status_code=403,
            detail="Administrators review reports but do not enter village data",
        )
    if user.role not in {"can_bo_thon", "to_cnscd"}:
        raise HTTPException(status_code=403, detail="Role cannot modify reports")
    if user.village_id and str(user.village_id) == str(village_id):
        return
    if user.role == "to_cnscd" and await _is_assigned_cnscd_village(
        repository,
        user,
        village_id,
    ):
        return
    raise HTTPException(status_code=403, detail="Cannot modify an unassigned village")


def _resolve_cnscd_assistance(
    user: UserProfile,
    requested: bool,
) -> tuple[bool, str | None]:
    """Derive assistance provenance from the authenticated profile only."""
    if not requested:
        return False, None
    if user.role != "to_cnscd":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Chỉ thành viên Tổ CNSCĐ mới được ghi nhận hỗ trợ nhập liệu.",
        )
    display_name = (user.display_name or "").strip()
    if not display_name:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail="Tài khoản Tổ CNSCĐ chưa có tên hiển thị hợp lệ.",
        )
    return True, display_name


async def _authorize_village_read(
    repository: ReportRepository,
    user: UserProfile,
    village_id: UUID,
) -> None:
    if user.role in {"admin_xa", "lanh_dao"}:
        return
    if user.role in {"can_bo_thon", "to_cnscd"} and str(user.village_id) == str(
        village_id
    ):
        return
    if user.role == "to_cnscd" and await _is_assigned_cnscd_village(
        repository,
        user,
        village_id,
    ):
        return
    raise HTTPException(status_code=403, detail="Cannot read an unassigned village")


def _sync_rejection(client_id: UUID, exc: HTTPException) -> RejectedReportItem:
    code = {
        403: "FORBIDDEN",
        404: "NOT_FOUND",
        409: "VERSION_CONFLICT",
        422: "VALIDATION_ERROR",
        502: "UPSTREAM_UNAVAILABLE",
        503: "UPSTREAM_UNAVAILABLE",
    }.get(exc.status_code, "SYNC_ERROR")
    detail = exc.detail
    if isinstance(detail, dict) and isinstance(detail.get("errors"), list):
        messages = [
            str(item.get("message"))
            for item in detail["errors"]
            if isinstance(item, dict) and item.get("message")
        ]
        message = "; ".join(messages) or "Report validation failed"
    elif isinstance(detail, str):
        message = detail
    else:
        message = "Unable to sync report"
    return RejectedReportItem(
        client_id=client_id,
        code=code,
        message=message,
        retryable=exc.status_code in {500, 502, 503, 504},
    )


def _known_report_values(values: dict[str, Any]) -> dict[str, Any]:
    return {code: values.get(code) for code in _indicator_codes()}


def _indicator_codes() -> list[str]:
    with RULES_PATH.open("r", encoding="utf-8") as rules_file:
        payload = json.load(rules_file)

    indicators = payload.get("indicators", [])
    if not isinstance(indicators, list):
        raise ValueError("validation_rules.json must contain an indicators list")

    return [
        str(indicator["code"])
        for indicator in indicators
        if isinstance(indicator, dict) and indicator.get("code")
    ]


def _format_export_value(value: Any) -> str:
    """Keep missing data blank; never misrepresent it as zero in exports."""
    return "" if value is None else _safe_document_text(value)


def _safe_document_text(value: Any) -> str:
    """Strip control characters that are illegal in XML/DOCX/PDF text."""
    return "".join(
        character
        for character in str(value)
        if character in "\t\n\r" or ord(character) >= 32
    )


def _export_matrix(
    reports_data: list[dict[str, Any]],
    villages_map: dict[str, str],
) -> list[list[str]]:
    """Build the authoritative CT01-CT14 matrix shared by DOCX/PDF/HTML."""
    indicator_codes = _indicator_codes()
    matrix = [["Thôn", *indicator_codes]]
    for report in reports_data:
        village_id = str(report.get("village_id", ""))
        village_name = villages_map.get(village_id, village_id)
        values = report.get("values")
        safe_values = values if isinstance(values, dict) else {}
        matrix.append(
            [
                _safe_document_text(village_name),
                *[
                    _format_export_value(safe_values.get(code))
                    for code in indicator_codes
                ],
            ]
        )
    return matrix


def _set_docx_run_font(run: Any, *, size: int = 8, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run_properties = run._element.get_or_add_rPr()
    run_properties.get_or_add_rFonts().set(qn("w:eastAsia"), "Times New Roman")


def _has_blocking_errors(errors: list[ValidationError]) -> bool:
    return any(error["error_type"] in BLOCKING_ERROR_TYPES for error in errors)


def _status_response(item: VillageSubmissionStatus) -> VillageStatusResponse:
    return VillageStatusResponse(
        village_id=UUID(item.village_id),
        village_name=item.village_name,
        old_village_names=item.old_village_names,
        report_id=UUID(item.report_id) if item.report_id is not None else None,
        submitted_at=item.submitted_at,
        due_date=item.due_date,
        days_late=item.days_late,
        status=item.status,
        dashboard_color=item.dashboard_color,
    )


# Export & Preview Helpers
async def safe_resolve_period(
    supabase: SupabaseAdminClient, period_id_or_name: str
) -> tuple[UUID, str]:
    try:
        uuid_val = UUID(period_id_or_name)
        rows = await supabase._rest_request(
            "GET", f"/rest/v1/report_periods?id=eq.{uuid_val}&select=id,name"
        )
        if rows:
            return UUID(rows[0]["id"]), str(rows[0]["name"])
    except ValueError:
        rows = await supabase._rest_request(
            "GET",
            f"/rest/v1/report_periods?name=eq.{quote(period_id_or_name, safe='')}&select=id,name",
        )
        if rows:
            return UUID(rows[0]["id"]), str(rows[0]["name"])

    raise HTTPException(
        status_code=404, detail=f"Không tìm thấy kỳ báo cáo '{period_id_or_name}'."
    )


async def resolve_period(
    supabase: SupabaseAdminClient, period_id_or_name: str
) -> tuple[str, str]:
    period_id, period_name = await safe_resolve_period(supabase, period_id_or_name)
    return str(period_id), period_name


async def get_villages_map(supabase: SupabaseAdminClient) -> dict[str, str]:
    rows = await supabase._rest_request("GET", "/rest/v1/villages?select=id,name")
    return {str(r["id"]): str(r["name"]) for r in rows}


async def get_period_reports_data(
    supabase: SupabaseAdminClient, period_id: str
) -> list[dict]:
    reports = await supabase._rest_request(
        "GET",
        (
            f"/rest/v1/reports?period_id=eq.{period_id}"
            # Timeliness records whether a report has ever been submitted. A
            # previously submitted report can move to needs_revision after an
            # approved citizen proposal and must remain part of the internal
            # snapshot/export while publication stays private.
            "&timeliness_status=in.(on_time,late)"
            "&select=id,village_id,workflow_status,timeliness_status,report_source,"
            "version,created_at,updated_at,submitted_at"
        ),
    )
    if not reports:
        return []

    report_ids = [r["id"] for r in reports]
    quoted_ids = ",".join(f'"{r_id}"' for r_id in report_ids)
    values = await supabase._rest_request(
        "GET",
        f"/rest/v1/report_values?report_id=in.({quoted_ids})&select=report_id,ct_code,value",
    )
    flags = await supabase._rest_request(
        "GET",
        (
            f"/rest/v1/report_validation_flags?report_id=in.({quoted_ids})"
            "&select=report_id,ct_code,error_type,message,resolved"
            "&order=created_at.asc"
        ),
    )

    indicator_codes = _indicator_codes()
    indicator_code_set = set(indicator_codes)
    vals_map: dict[str, dict[str, int | None]] = {}
    for val in values:
        code = str(val.get("ct_code", ""))
        if code not in indicator_code_set:
            continue
        r_id = str(val["report_id"])
        if r_id not in vals_map:
            vals_map[r_id] = {}
        raw_value = val.get("value")
        vals_map[r_id][code] = int(raw_value) if raw_value is not None else None

    result = []
    flags_map: dict[str, list[dict[str, Any]]] = {}
    for flag in flags:
        flags_map.setdefault(str(flag["report_id"]), []).append(
            {
                "ct_code": str(flag.get("ct_code") or ""),
                "error_type": str(flag.get("error_type") or ""),
                "message": str(flag.get("message") or ""),
                "resolved": bool(flag.get("resolved", False)),
            }
        )
    for raw_report in reports:
        report = dict(raw_report)
        r_vals = vals_map.get(str(report["id"]), {})
        for code in indicator_codes:
            if code not in r_vals:
                r_vals[code] = None
        report["values"] = r_vals
        report["validation_flags"] = flags_map.get(str(report["id"]), [])
        report["rule_version"] = "2026-07"
        result.append(report)

    return result


def generate_docx_file(
    period_name: str,
    reports_data: list,
    villages_map: dict,
    scope_name: str | None = None,
) -> bytes:
    doc = docx.Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Inches(0.35)
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(8)
    normal_fonts = normal_style.element.get_or_add_rPr().get_or_add_rFonts()
    normal_fonts.set(qn("w:eastAsia"), "Times New Roman")
    heading_style = doc.styles["Heading 1"]
    heading_style.font.name = "Times New Roman"
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    heading_style.element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), "Times New Roman"
    )

    doc.add_heading("BÁO CÁO VĂN HÓA - XÃ HỘI XÃ BÀ NÀ", level=1)
    doc.add_paragraph(f"Kỳ báo cáo: {_safe_document_text(period_name)}")
    if scope_name:
        doc.add_paragraph(f"Phạm vi: {_safe_document_text(scope_name)}")

    if not reports_data:
        doc.add_paragraph("Lưu ý: Chưa có dữ liệu báo cáo cho kỳ này.")
    else:
        matrix = _export_matrix(reports_data, villages_map)
        table = doc.add_table(rows=0, cols=len(matrix[0]))
        table.style = "Light Shading Accent 1"
        table.autofit = False
        column_widths = [Inches(1.8), *[Inches(0.60) for _ in _indicator_codes()]]
        for row_index, matrix_row in enumerate(matrix):
            row_cells = table.add_row().cells
            for column_index, cell_text in enumerate(matrix_row):
                cell = row_cells[column_index]
                cell.width = column_widths[column_index]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_docx_run_font(run, bold=row_index == 0)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_pdf_file(
    period_name: str, reports_data: list, villages_map: dict
) -> bytes:
    from matplotlib import font_manager
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    buffer = io.BytesIO()
    regular_font = font_manager.findfont("DejaVu Sans")
    bold_font = font_manager.findfont(
        font_manager.FontProperties(family="DejaVu Sans", weight="bold")
    )
    registered_fonts = set(pdfmetrics.getRegisteredFontNames())
    if "BaNaUnicode" not in registered_fonts:
        pdfmetrics.registerFont(TTFont("BaNaUnicode", regular_font))
    if "BaNaUnicode-Bold" not in registered_fonts:
        pdfmetrics.registerFont(TTFont("BaNaUnicode-Bold", bold_font))
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        rightMargin=20,
        leftMargin=20,
        topMargin=24,
        bottomMargin=24,
    )
    story = []

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Heading1"],
        fontName="BaNaUnicode-Bold",
        fontSize=18,
        textColor=colors.HexColor("#1E293B"),
        spaceAfter=15,
        alignment=1,
    )

    body_style = ParagraphStyle(
        "BodyTextCustom",
        parent=styles["BodyText"],
        fontName="BaNaUnicode",
        fontSize=10,
        textColor=colors.HexColor("#334155"),
        spaceAfter=8,
    )

    story.append(Paragraph("BÁO CÁO VĂN HÓA - XÃ HỘI XÃ BÀ NÀ", title_style))
    story.append(
        Paragraph(
            f"Kỳ báo cáo: {html.escape(_safe_document_text(period_name))}",
            ParagraphStyle("Sub", parent=title_style, fontSize=12, spaceAfter=20),
        )
    )
    story.append(Spacer(1, 10))

    if not reports_data:
        story.append(
            Paragraph("Lưu ý: Chưa có dữ liệu báo cáo cho kỳ này.", body_style)
        )
    else:
        matrix = _export_matrix(reports_data, villages_map)
        header_cell_style = ParagraphStyle(
            "ExportHeaderCell",
            fontName="BaNaUnicode-Bold",
            fontSize=6,
            leading=7,
            alignment=1,
        )
        body_cell_style = ParagraphStyle(
            "ExportBodyCell",
            fontName="BaNaUnicode",
            fontSize=6,
            leading=7,
            alignment=1,
        )
        table_data = [
            [
                Paragraph(
                    html.escape(cell),
                    header_cell_style if row_index == 0 else body_cell_style,
                )
                for cell in row
            ]
            for row_index, row in enumerate(matrix)
        ]
        usable_width = landscape(A4)[0] - 40
        village_width = 105
        indicator_width = (usable_width - village_width) / len(_indicator_codes())
        t = Table(
            table_data,
            colWidths=[village_width, *[indicator_width for _ in _indicator_codes()]],
            repeatRows=1,
            splitByRow=1,
        )
        t.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F1F5F9")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0F172A")),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("FONTNAME", (0, 0), (-1, 0), "BaNaUnicode-Bold"),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                    ("FONTNAME", (0, 1), (-1, -1), "BaNaUnicode"),
                    ("FONTSIZE", (0, 0), (-1, -1), 6),
                    ("ALIGN", (0, 1), (0, -1), "LEFT"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 3),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        story.append(t)

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_preview_html(
    period_name: str, reports_data: list, villages_map: dict
) -> str:
    indicator_codes = _indicator_codes()
    table_headers_html = "".join(
        f"<th>{html.escape(code)}</th>" for code in indicator_codes
    )
    if not reports_data:
        table_rows_html = f"""
        <tr>
            <td colspan="{len(indicator_codes) + 1}" style="text-align: center; padding: 30px; color: #64748b; font-style: italic;">
                Lưu ý: Chưa có dữ liệu báo cáo cho kỳ này.
            </td>
        </tr>
        """
    else:
        matrix = _export_matrix(reports_data, villages_map)
        table_rows_html = "".join(
            '<tr style="border-bottom: 1px solid #e2e8f0;">'
            f'<td style="padding: 12px 16px; font-weight: 500; color: #0f172a;">{html.escape(row[0])}</td>'
            + "".join(
                f'<td style="padding: 12px 16px; text-align: center;">{html.escape(value)}</td>'
                for value in row[1:]
            )
            + "</tr>"
            for row in matrix[1:]
        )

    return f"""
    <!DOCTYPE html>
    <html lang="vi">
    <head>
        <meta charset="UTF-8">
        <title>Xem trước báo cáo - Ba Na SmartLink</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                background-color: #f8fafc;
                margin: 0;
                padding: 40px 20px;
                color: #334155;
            }}
            .container {{
                max-width: 1400px;
                margin: 0 auto;
                background: white;
                border-radius: 16px;
                box-shadow: 0 4px 6px -1px rgb(0 0 0 / 0.1), 0 2px 4px -2px rgb(0 0 0 / 0.1);
                padding: 40px;
                border: 1px solid #e2e8f0;
            }}
            .header {{
                text-align: center;
                margin-bottom: 30px;
                border-bottom: 2px solid #f1f5f9;
                padding-bottom: 20px;
            }}
            .header h1 {{
                color: #0f172a;
                font-size: 24px;
                margin: 0 0 10px 0;
                font-weight: 700;
                letter-spacing: -0.025em;
            }}
            .header p {{
                color: #64748b;
                margin: 0;
                font-size: 14px;
            }}
            .badge {{
                display: inline-block;
                padding: 4px 12px;
                background-color: #3b82f6;
                color: white;
                border-radius: 9999px;
                font-weight: 500;
                font-size: 12px;
                margin-top: 8px;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-top: 20px;
            }}
            .table-wrap {{ overflow-x: auto; }}
            th {{
                background-color: #f1f5f9;
                color: #475569;
                font-weight: 600;
                text-align: center;
                padding: 12px 16px;
                font-size: 13px;
                text-transform: uppercase;
                letter-spacing: 0.05em;
            }}
            th:first-child {{
                text-align: left;
                border-top-left-radius: 8px;
                border-bottom-left-radius: 8px;
            }}
            th:last-child {{
                border-top-right-radius: 8px;
                border-bottom-right-radius: 8px;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1>BÁO CÁO VĂN HÓA - XÃ HỘI XÃ BÀ NÀ</h1>
                <p>Hệ thống Quản lý và Số hóa Dữ liệu Đồng bộ Ba Na SmartLink</p>
                <span class="badge">Kỳ báo cáo: {html.escape(_safe_document_text(period_name))}</span>
            </div>
            <div class="table-wrap"><table>
                <thead>
                    <tr>
                        <th>Thôn</th>
                        {table_headers_html}
                    </tr>
                </thead>
                <tbody>
                    {table_rows_html}
                </tbody>
            </table></div>
        </div>
    </body>
    </html>
    """


class ApproveReportRequest(BaseModel):
    action: Literal["approve", "lock"]
    expected_version: int = Field(ge=1)


@router.patch("/{report_id}/approve")
async def approve_or_lock_report(
    report_id: UUID,
    payload: ApproveReportRequest,
    repository: Annotated[ReportRepository, Depends(get_report_repository)],
    admin: Annotated[UserProfile, Depends(require_admin_xa)],
) -> dict:
    """Approve or lock a report as admin_xa using optimistic locking."""
    new_status = "approved" if payload.action == "approve" else "locked"
    _ = admin
    try:
        rows = await repository._supabase._rest_request(
            "POST",
            "/rest/v1/rpc/transition_report_workflow",
            {
                "p_report_id": str(report_id),
                "p_expected_version": payload.expected_version,
                "p_action": payload.action,
            },
        )
    except SupabaseAdminError as exc:
        if exc.error_code == "42501" or exc.status_code == 403:
            raise HTTPException(
                status_code=403, detail="Only admin can approve reports"
            ) from exc
        if exc.error_code in {"22023", "40001"} or exc.status_code in {400, 409}:
            raise HTTPException(
                status_code=409,
                detail="Report version or state conflict",
            ) from exc
        raise HTTPException(status_code=502, detail="Unable to update report") from exc

    if not rows:
        raise HTTPException(status_code=409, detail="Report version or state conflict")

    return {
        "report_id": str(report_id),
        "workflow_status": new_status,
        "version": int(rows[0]["version"]),
    }


@router.patch("/{report_id}/publish")
async def publish_report(
    report_id: UUID,
    expected_version: int,
    repository: Annotated[ReportRepository, Depends(get_report_repository)],
    admin: Annotated[UserProfile, Depends(require_admin_xa)],
) -> dict[str, Any]:
    _ = admin
    try:
        rows = await repository._supabase._rest_request(
            "POST",
            "/rest/v1/rpc/transition_report_workflow",
            {
                "p_report_id": str(report_id),
                "p_expected_version": expected_version,
                "p_action": "publish",
            },
        )
    except SupabaseAdminError as exc:
        if exc.error_code == "42501" or exc.status_code == 403:
            raise HTTPException(
                status_code=403, detail="Only admin can publish reports"
            ) from exc
        if exc.error_code in {"22023", "40001"} or exc.status_code in {400, 409}:
            raise HTTPException(
                status_code=409,
                detail="Report must be approved and unchanged",
            ) from exc
        raise HTTPException(status_code=502, detail="Unable to publish report") from exc
    if not rows:
        raise HTTPException(
            status_code=409, detail="Report must be approved and unchanged"
        )
    return {
        "report_id": str(report_id),
        "publication_status": "published",
        "version": int(rows[0]["version"]),
    }


@router.get("/export/{file_format}")
async def export_reports(
    file_format: str,
    period_id: str,
    repository: Annotated[ReportRepository, Depends(get_report_repository)],
    current_user: Annotated[UserProfile, Depends(require_admin_or_leader)],
):
    supabase = repository._supabase
    period_uuid, period_name = await resolve_period(supabase, period_id)
    villages_map = await get_villages_map(supabase)
    reports_data = await get_period_reports_data(supabase, period_uuid)

    if file_format == "xlsx":
        file_bytes = generate_summary_xlsx_file(period_name, reports_data, villages_map)
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        filename = f"Bang_tong_hop_Bana_SmartLink_{period_name}.xlsx"
    elif file_format == "docx":
        file_bytes = generate_docx_file(period_name, reports_data, villages_map)
        media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        filename = f"Bao_cao_Bana_SmartLink_{period_name}.docx"
    elif file_format == "pdf":
        file_bytes = generate_pdf_file(period_name, reports_data, villages_map)
        media_type = "application/pdf"
        filename = f"Bao_cao_Bana_SmartLink_{period_name}.pdf"
    else:
        raise HTTPException(
            status_code=400,
            detail="Định dạng xuất bản không hỗ trợ. Chỉ hỗ trợ xlsx, docx, pdf.",
        )

    return Response(
        content=file_bytes,
        media_type=media_type,
        headers={"Content-Disposition": _content_disposition(filename)},
    )


@router.get("/village/{village_id}/export/{file_format}")
async def export_village_report(
    village_id: UUID,
    file_format: str,
    period_id: str,
    repository: Annotated[ReportRepository, Depends(get_report_repository)],
    current_user: Annotated[UserProfile, Depends(require_authenticated_user)],
):
    await _authorize_village_read(repository, current_user, village_id)
    supabase = repository._supabase
    if file_format not in {"xlsx", "docx"}:
        raise HTTPException(
            status_code=400,
            detail="Báo cáo theo thôn chỉ hỗ trợ định dạng XLSX hoặc DOCX.",
        )

    period_uuid, period_name = await resolve_period(supabase, period_id)
    villages_map = await get_villages_map(supabase)
    village_id_text = str(village_id)
    village_name = villages_map.get(village_id_text, f"Thôn {village_id_text}")

    reports_data = await get_period_reports_data(supabase, period_uuid)

    # Find the report for this village
    village_report = next(
        (r for r in reports_data if r["village_id"] == village_id_text), None
    )
    if not village_report:
        raise HTTPException(
            status_code=404, detail="Không tìm thấy báo cáo của thôn trong kỳ này."
        )

    if file_format == "xlsx":
        file_bytes = generate_village_xlsx_file(
            period_name, village_report, village_name
        )
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        filename = f"Phieu_bao_cao_{village_name}_{period_name}.xlsx".replace(" ", "_")
    else:
        file_bytes = generate_docx_file(
            period_name,
            [village_report],
            {village_id_text: village_name},
            scope_name=village_name,
        )
        media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        filename = f"Phieu_bao_cao_{village_name}_{period_name}.docx".replace(" ", "_")

    return Response(
        content=file_bytes,
        media_type=media_type,
        headers={"Content-Disposition": _content_disposition(filename)},
    )


@router.get("/preview/{file_format}", response_class=HTMLResponse)
async def preview_reports(
    file_format: str,
    period_id: str,
    repository: Annotated[ReportRepository, Depends(get_report_repository)],
    current_user: Annotated[UserProfile, Depends(require_authenticated_user)],
):
    _ = current_user
    supabase = repository._supabase
    period_uuid, period_name = await resolve_period(supabase, period_id)
    villages_map = await get_villages_map(supabase)
    reports_data = await get_period_reports_data(supabase, period_uuid)

    html_content = generate_preview_html(period_name, reports_data, villages_map)
    return HTMLResponse(content=html_content)


def _content_disposition(filename: str) -> str:
    safe_ascii = "".join(
        character
        if character.isascii() and (character.isalnum() or character in "._-")
        else "_"
        for character in filename
    )
    return f"attachment; filename=\"{safe_ascii}\"; filename*=UTF-8''{quote(filename, safe='')}"


__all__ = ["period_router", "router"]
