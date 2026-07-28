from __future__ import annotations

import html
import io
from datetime import datetime
from typing import Any

import docx
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from services.export_service import (
    INDICATORS_DICT,
    _format_submitted_at,
    _ordered_village_items,
    _source_label,
    _submission_status_label,
    _timeliness_label,
    _workflow_label,
    build_semantic_summary,
    semantic_coverage_text,
    semantic_provenance_text,
)
from services.metric_registry import load_metric_registry


AUTHORITY_NAME = "ỦY BAN NHÂN DÂN\nXÃ BÀ NÀ"
NATIONAL_NAME = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
NATIONAL_MOTTO = "Độc lập - Tự do - Hạnh phúc"
BODY_FONT = "Times New Roman"
TABLE_HEADER_COLOR = "1F4E78"
TABLE_TOTAL_COLOR = "D9EAF7"
TABLE_ALT_COLOR = "F4F8FB"
INDICATOR_GROUPS = (
    ("Nhóm 1: Quy mô dân cư và an sinh", ("CT01", "CT02", "CT03", "CT04")),
    ("Nhóm 2: Chính sách xã hội và trẻ em", ("CT05", "CT06", "CT07", "CT08")),
    ("Nhóm 3: Văn hóa, lao động và y tế", ("CT09", "CT10", "CT11", "CT12")),
    ("Nhóm 4: Dịch vụ công và an toàn xã hội", ("CT13", "CT14")),
)


def _safe_text(value: Any) -> str:
    return "".join(
        character
        for character in str(value if value is not None else "")
        if character in "\t\n\r" or ord(character) >= 32
    )


def _format_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return _safe_text(value)
    if isinstance(value, int):
        return f"{value:,}".replace(",", ".")
    if isinstance(value, float):
        return f"{value:,.2f}".rstrip("0").rstrip(".").replace(",", ".")
    return _safe_text(value)


def _report_index(reports_data: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    return {
        str(report.get("village_id")): report
        for report in reports_data
        if report.get("village_id") is not None
    }


def _submitted_totals(
    reports_data: list[dict[str, Any]],
) -> tuple[dict[str, int | float], set[str]]:
    totals: dict[str, int | float] = {code: 0 for code in INDICATORS_DICT}
    incomplete: set[str] = set()
    for report in reports_data:
        values = report.get("values") if isinstance(report.get("values"), dict) else {}
        for code in INDICATORS_DICT:
            value = values.get(code)
            if isinstance(value, bool) or not isinstance(value, (int, float)):
                incomplete.add(code)
            else:
                totals[code] += value
    return totals, incomplete


def _semantic_summary_rows(
    *,
    period_name: str,
    reports_data: list[dict[str, Any]],
    villages_map: dict[str, str],
    period_id: str | None,
) -> list[list[str]]:
    registry = load_metric_registry()
    rows = [["Chỉ số", "Giá trị", "Đơn vị", "Độ phủ", "Căn cứ"]]
    for result in build_semantic_summary(
        period_name,
        reports_data,
        villages_map,
        period_id=period_id,
    ):
        definition = registry.get(result.metric_id)
        if result.value is None:
            value = "—"
        elif result.unit and result.unit.startswith("percent_"):
            value = f"{result.value:.1f}%".replace(".", ",")
        else:
            value = _format_value(result.value)
        rows.append(
            [
                definition.label_vi if definition else result.metric_id,
                value,
                (
                    definition.display_unit_vi
                    if definition
                    else str(result.unit or "")
                ),
                semantic_coverage_text(result),
                semantic_provenance_text(result),
            ]
        )
    return rows


def _set_docx_run(
    run: Any,
    *,
    size: float = 12,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run_properties = run._element.get_or_add_rPr()
    fonts = run_properties.get_or_add_rFonts()
    for font_key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(font_key), BODY_FONT)


def _format_docx_paragraph(
    paragraph: Any,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    before: float = 0,
    after: float = 0,
    line_spacing: float = 1.15,
) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def _set_docx_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_docx_cell_text(
    cell: Any,
    value: Any,
    *,
    size: float = 12,
    bold: bool = False,
    color: str | None = None,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    _format_docx_paragraph(paragraph, alignment=alignment, after=0)
    run = paragraph.add_run(_safe_text(value))
    _set_docx_run(run, size=size, bold=bold)
    if color:
        run.font.color.rgb = docx.shared.RGBColor.from_string(color)


def _repeat_docx_table_header(row: Any) -> None:
    row_properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    row_properties.append(header)


def _keep_docx_row_together(row: Any) -> None:
    row_properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    row_properties.append(cant_split)


def _add_docx_table(
    document: Any,
    rows: list[list[Any]],
    *,
    header_rows: int = 1,
    first_column_left: bool = False,
    font_size: float = 12,
) -> Any:
    table = document.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row_values in enumerate(rows):
        row = table.add_row()
        _keep_docx_row_together(row)
        for column_index, value in enumerate(row_values):
            is_header = row_index < header_rows
            is_total = row_index == len(rows) - 1 and _safe_text(row_values[0]) in {
                "",
                "TỔNG CỘNG",
            }
            _set_docx_cell_text(
                row.cells[column_index],
                value,
                size=font_size,
                bold=is_header or is_total,
                color="FFFFFF" if is_header else None,
                alignment=(
                    WD_ALIGN_PARAGRAPH.LEFT
                    if first_column_left and column_index == 1 and not is_header
                    else WD_ALIGN_PARAGRAPH.CENTER
                ),
            )
            if is_header:
                _set_docx_cell_shading(row.cells[column_index], TABLE_HEADER_COLOR)
            elif is_total:
                _set_docx_cell_shading(row.cells[column_index], TABLE_TOTAL_COLOR)
            elif row_index % 2 == 0:
                _set_docx_cell_shading(row.cells[column_index], TABLE_ALT_COLOR)
        if row_index < header_rows:
            _repeat_docx_table_header(row)
    document.add_paragraph()
    return table


def _add_docx_administrative_header(document: Any, *, village_name: str | None) -> None:
    header = document.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    left, right = header.rows[0].cells
    authority = (
        f"{AUTHORITY_NAME}\n{_safe_text(village_name).upper()}"
        if village_name
        else AUTHORITY_NAME
    )
    _set_docx_cell_text(
        left,
        authority,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    right.text = ""
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right_paragraph = right.paragraphs[0]
    _format_docx_paragraph(
        right_paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        after=0,
    )
    _set_docx_run(right_paragraph.add_run(NATIONAL_NAME), bold=True)
    right_paragraph.add_run("\n")
    motto_run = right_paragraph.add_run(NATIONAL_MOTTO)
    _set_docx_run(motto_run, bold=True)
    motto_run.font.underline = True
    for cell in header.rows[0].cells:
        properties = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "nil")
            borders.append(element)
        properties.append(borders)


def _add_docx_title(
    document: Any,
    *,
    period_name: str,
    village_name: str | None,
) -> None:
    title = (
        "PHIẾU BÁO CÁO SỐ LIỆU VĂN HÓA – XÃ HỘI ĐỊNH KỲ"
        if village_name
        else "BÁO CÁO TỔNG HỢP SỐ LIỆU VĂN HÓA – XÃ HỘI"
    )
    paragraph = document.add_paragraph()
    _format_docx_paragraph(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        before=12,
        after=4,
    )
    _set_docx_run(paragraph.add_run(title), size=14, bold=True)
    period_paragraph = document.add_paragraph()
    _format_docx_paragraph(
        period_paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        after=10,
    )
    _set_docx_run(
        period_paragraph.add_run(f"Kỳ báo cáo: {_safe_text(period_name)}"),
        size=12,
        bold=True,
    )


def _add_docx_heading(document: Any, text: str) -> None:
    paragraph = document.add_paragraph()
    _format_docx_paragraph(paragraph, before=8, after=4)
    _set_docx_run(paragraph.add_run(text), size=12, bold=True)


def _add_docx_page_numbers(document: Any) -> None:
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        _format_docx_paragraph(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            after=0,
        )
        _set_docx_run(paragraph.add_run("Trang "), size=11)
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        paragraph._p.append(field)


def _add_docx_signature(document: Any) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(
        table.rows[0].cells,
        (
            "NGƯỜI LẬP BIỂU\n(Ký, ghi rõ họ tên)",
            "TM. ỦY BAN NHÂN DÂN XÃ\n(Ký số hoặc ký, đóng dấu)",
        ),
        strict=True,
    ):
        _set_docx_cell_text(
            cell,
            f"{text}\n\n\n\n",
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )


def _docx_village_body(
    document: Any,
    *,
    report: dict[str, Any] | None,
    village_name: str,
) -> None:
    report = report or {}
    metadata_rows = [
        ["Đơn vị báo cáo", village_name],
        ["Người lập báo cáo", report.get("reporter_name") or ""],
        ["Chức danh", report.get("reporter_title") or "Cán bộ thôn"],
        ["Số điện thoại", report.get("reporter_phone") or ""],
        ["Thời điểm nộp", _format_submitted_at(report.get("submitted_at"))],
        ["Trạng thái", _submission_status_label(report) if report else "Chưa nộp"],
    ]
    _add_docx_table(
        document,
        [["Thông tin", "Nội dung"], *metadata_rows],
        first_column_left=True,
    )
    _add_docx_heading(document, "I. SỐ LIỆU BÁO CÁO")
    values = report.get("values") if isinstance(report.get("values"), dict) else {}
    rows = [["Mã CT", "Tên chỉ tiêu", "Đơn vị tính", "Số liệu", "Ghi chú"]]
    for code, (name, unit) in INDICATORS_DICT.items():
        rows.append([code, name, unit, _format_value(values.get(code)), ""])
    _add_docx_table(document, rows, first_column_left=True)
    _add_docx_heading(document, "II. XÁC NHẬN")
    paragraph = document.add_paragraph(
        "Người lập biểu chịu trách nhiệm rà soát số liệu, đối chiếu hồ sơ nguồn "
        "và xác nhận nội dung trước khi gửi phê duyệt."
    )
    _format_docx_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=12)
    for run in paragraph.runs:
        _set_docx_run(run)
    _add_docx_signature(document)


def _docx_summary_body(
    document: Any,
    *,
    period_name: str,
    reports_data: list[dict[str, Any]],
    villages_map: dict[str, str],
    period_id: str | None,
) -> None:
    report_by_village = _report_index(reports_data)
    villages = _ordered_village_items(villages_map)
    on_time = sum(
        1 for report in reports_data if report.get("timeliness_status") == "on_time"
    )
    late = sum(
        1 for report in reports_data if report.get("timeliness_status") == "late"
    )
    summary_rows = [
        ["Tổng số thôn", "Đã nộp", "Đúng hạn", "Trễ hạn", "Chưa nộp", "Tỷ lệ nộp"],
        [
            len(villages),
            len(reports_data),
            on_time,
            late,
            max(0, len(villages) - len(reports_data)),
            f"{(len(reports_data) * 100 / len(villages)):.1f}%"
            if villages
            else "0%",
        ],
    ]
    _add_docx_heading(document, "I. TÌNH HÌNH NỘP BÁO CÁO")
    _add_docx_table(document, summary_rows)
    progress_rows = [["STT", "Thôn", "Thời điểm nộp", "Trạng thái"]]
    for index, (village_id, village_name) in enumerate(villages, start=1):
        report = report_by_village.get(str(village_id))
        progress_rows.append(
            [
                index,
                village_name,
                _format_submitted_at(report.get("submitted_at")) if report else "",
                _submission_status_label(report) if report else "Chưa nộp",
            ]
        )
    _add_docx_table(document, progress_rows, first_column_left=True)

    _add_docx_heading(
        document,
        "II. CHỈ SỐ NGỮ NGHĨA — HỒ SƠ ĐÃ DUYỆT/KHÓA",
    )
    _add_docx_table(
        document,
        _semantic_summary_rows(
            period_name=period_name,
            reports_data=reports_data,
            villages_map=villages_map,
            period_id=period_id,
        ),
        first_column_left=True,
    )

    totals, incomplete = _submitted_totals(reports_data)
    document.add_page_break()
    _add_docx_heading(document, "III. BẢNG TỔNG HỢP CHỈ TIÊU THEO THÔN")
    for group_index, (group_name, codes) in enumerate(INDICATOR_GROUPS):
        if group_index:
            document.add_page_break()
        paragraph = document.add_paragraph()
        _format_docx_paragraph(paragraph, after=3)
        _set_docx_run(paragraph.add_run(group_name), size=12, bold=True, italic=True)
        rows: list[list[Any]] = [["STT", "Thôn", *codes]]
        for index, (village_id, village_name) in enumerate(villages, start=1):
            report = report_by_village.get(str(village_id))
            values = (
                report.get("values")
                if report and isinstance(report.get("values"), dict)
                else {}
            )
            rows.append(
                [
                    index,
                    village_name,
                    *[_format_value(values.get(code)) for code in codes],
                ]
            )
        rows.append(
            [
                "",
                "TỔNG CỘNG",
                *[
                    "—" if code in incomplete else _format_value(totals[code])
                    for code in codes
                ],
            ]
        )
        _add_docx_table(document, rows, first_column_left=True)

    document.add_page_break()
    _add_docx_heading(document, "IV. DANH MỤC CHỈ TIÊU")
    dictionary_rows = [["Mã CT", "Tên chỉ tiêu", "Đơn vị tính"]]
    dictionary_rows.extend(
        [code, name, unit] for code, (name, unit) in INDICATORS_DICT.items()
    )
    _add_docx_table(document, dictionary_rows, first_column_left=True)

    document.add_page_break()
    _add_docx_heading(document, "V. NGUỒN DỮ LIỆU VÀ CẢNH BÁO")
    source_rows = [
        ["Thôn", "Nguồn nhập", "Trạng thái", "Đúng hạn", "Thời điểm nộp"]
    ]
    for village_id, village_name in villages:
        report = report_by_village.get(str(village_id))
        source_rows.append(
            [
                village_name,
                _source_label(report.get("report_source")) if report else "",
                _workflow_label(report.get("workflow_status")) if report else "Chưa nộp",
                _timeliness_label(report.get("timeliness_status")) if report else "",
                _format_submitted_at(report.get("submitted_at")) if report else "",
            ]
        )
    _add_docx_table(document, source_rows, first_column_left=True)
    warning_rows = [["Thôn", "Chỉ tiêu", "Loại cảnh báo", "Nội dung", "Trạng thái"]]
    for report in reports_data:
        village_name = villages_map.get(
            str(report.get("village_id")), str(report.get("village_id") or "")
        )
        for flag in report.get("validation_flags") or []:
            warning_rows.append(
                [
                    village_name,
                    flag.get("ct_code") or "",
                    flag.get("error_type") or "",
                    flag.get("message") or "",
                    "Đã xử lý" if flag.get("resolved") else "Cần xem",
                ]
            )
    if len(warning_rows) == 1:
        warning_rows.append(
            ["", "", "", "Không có cảnh báo dữ liệu trong phạm vi xuất báo cáo.", ""]
        )
    _add_docx_table(document, warning_rows, first_column_left=True)
    _add_docx_signature(document)


def generate_docx_file(
    period_name: str,
    reports_data: list,
    villages_map: dict,
    scope_name: str | None = None,
    *,
    period_id: str | None = None,
) -> bytes:
    """Generate an A4 administrative report with readable, repeatable tables."""
    document = docx.Document()
    section = document.sections[0]
    if scope_name is None:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(18)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    normal_style = document.styles["Normal"]
    normal_style.font.name = BODY_FONT
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_fonts = normal_style.element.get_or_add_rPr().get_or_add_rFonts()
    for font_key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        normal_fonts.set(qn(font_key), BODY_FONT)

    document.core_properties.title = (
        "Phiếu báo cáo số liệu văn hóa – xã hội"
        if scope_name
        else "Báo cáo tổng hợp số liệu văn hóa – xã hội xã Bà Nà"
    )
    document.core_properties.subject = f"Kỳ báo cáo: {_safe_text(period_name)}"
    document.core_properties.author = "Ủy ban nhân dân xã Bà Nà"
    document.core_properties.keywords = "báo cáo hành chính; Bà Nà; văn hóa xã hội"

    _add_docx_administrative_header(document, village_name=scope_name)
    _add_docx_title(
        document,
        period_name=period_name,
        village_name=scope_name,
    )
    context = document.add_paragraph(
        f"Phạm vi: {_safe_text(scope_name or 'Toàn xã')}  |  "
        f"Thời điểm kết xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )
    _format_docx_paragraph(
        context,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        after=10,
    )
    for run in context.runs:
        _set_docx_run(run, size=11, italic=True)

    if scope_name:
        report = reports_data[0] if reports_data else None
        _docx_village_body(
            document,
            report=report,
            village_name=scope_name,
        )
    else:
        _docx_summary_body(
            document,
            period_name=period_name,
            reports_data=reports_data,
            villages_map=villages_map,
            period_id=period_id,
        )
    _add_docx_page_numbers(document)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _register_pdf_fonts() -> tuple[str, str, str]:
    from matplotlib import get_data_path
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_dir = f"{get_data_path()}/fonts/ttf"
    names = ("BaNaTimes", "BaNaTimes-Bold", "BaNaTimes-Italic")
    paths = (
        f"{font_dir}/DejaVuSerif.ttf",
        f"{font_dir}/DejaVuSerif-Bold.ttf",
        f"{font_dir}/DejaVuSerif-Italic.ttf",
    )
    registered = set(pdfmetrics.getRegisteredFontNames())
    for name, path in zip(names, paths, strict=True):
        if name not in registered:
            pdfmetrics.registerFont(TTFont(name, path))
    return names


def _pdf_paragraph(value: Any, style: ParagraphStyle) -> Paragraph:
    return Paragraph(html.escape(_safe_text(value)).replace("\n", "<br/>"), style)


def _pdf_table(
    rows: list[list[Any]],
    *,
    styles: dict[str, ParagraphStyle],
    column_widths: list[float] | None = None,
    first_column_left: bool = False,
) -> Table:
    data: list[list[Paragraph]] = []
    for row_index, row in enumerate(rows):
        data.append(
            [
                _pdf_paragraph(
                    value,
                    styles["table_header"] if row_index == 0 else styles["table"],
                )
                for value in row
            ]
        )
    table = Table(
        data,
        colWidths=column_widths,
        repeatRows=1,
        splitByRow=1,
        hAlign="CENTER",
    )
    commands: list[tuple[Any, ...]] = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{TABLE_HEADER_COLOR}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "BaNaTimes-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#64748B")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]
    if first_column_left:
        commands.append(("ALIGN", (1, 1), (1, -1), "LEFT"))
    for row_index in range(2, len(rows), 2):
        commands.append(
            (
                "BACKGROUND",
                (0, row_index),
                (-1, row_index),
                colors.HexColor(f"#{TABLE_ALT_COLOR}"),
            )
        )
    if rows and _safe_text(rows[-1][0]) in {"", "TỔNG CỘNG"}:
        commands.extend(
            [
                (
                    "BACKGROUND",
                    (0, -1),
                    (-1, -1),
                    colors.HexColor(f"#{TABLE_TOTAL_COLOR}"),
                ),
                ("FONTNAME", (0, -1), (-1, -1), "BaNaTimes-Bold"),
            ]
        )
    table.setStyle(TableStyle(commands))
    return table


def _pdf_styles() -> dict[str, ParagraphStyle]:
    _, bold, italic = _register_pdf_fonts()
    base = getSampleStyleSheet()
    return {
        "body": ParagraphStyle(
            "AdministrativeBody",
            parent=base["BodyText"],
            fontName="BaNaTimes",
            fontSize=12,
            leading=13.8,
            textColor=colors.black,
            spaceAfter=6,
            alignment=TA_LEFT,
        ),
        "small": ParagraphStyle(
            "AdministrativeSmall",
            parent=base["BodyText"],
            fontName=italic,
            fontSize=10,
            leading=11.5,
            textColor=colors.HexColor("#334155"),
            alignment=TA_CENTER,
        ),
        "header": ParagraphStyle(
            "AdministrativeHeader",
            parent=base["BodyText"],
            fontName=bold,
            fontSize=11,
            leading=14,
            alignment=TA_CENTER,
        ),
        "title": ParagraphStyle(
            "AdministrativeTitle",
            parent=base["Heading1"],
            fontName=bold,
            fontSize=14,
            leading=17,
            alignment=TA_CENTER,
            textColor=colors.black,
            spaceBefore=10,
            spaceAfter=6,
        ),
        "heading": ParagraphStyle(
            "AdministrativeHeading",
            parent=base["Heading2"],
            fontName=bold,
            fontSize=12,
            leading=14,
            textColor=colors.black,
            spaceBefore=8,
            spaceAfter=5,
        ),
        "subheading": ParagraphStyle(
            "AdministrativeSubheading",
            parent=base["BodyText"],
            fontName=italic,
            fontSize=11,
            leading=13,
            textColor=colors.black,
            spaceBefore=5,
            spaceAfter=4,
        ),
        "table_header": ParagraphStyle(
            "AdministrativeTableHeader",
            parent=base["BodyText"],
            fontName=bold,
            fontSize=9,
            leading=10,
            textColor=colors.white,
            alignment=TA_CENTER,
        ),
        "table": ParagraphStyle(
            "AdministrativeTable",
            parent=base["BodyText"],
            fontName="BaNaTimes",
            fontSize=9,
            leading=10,
            textColor=colors.black,
            alignment=TA_CENTER,
        ),
    }


def _pdf_administrative_header(
    *,
    styles: dict[str, ParagraphStyle],
    village_name: str | None,
) -> list[Any]:
    authority = (
        f"{AUTHORITY_NAME}\n{_safe_text(village_name).upper()}"
        if village_name
        else AUTHORITY_NAME
    )
    header = Table(
        [
            [
                _pdf_paragraph(authority, styles["header"]),
                Paragraph(
                    f"{html.escape(NATIONAL_NAME)}<br/>"
                    f"<u>{html.escape(NATIONAL_MOTTO)}</u>",
                    styles["header"],
                ),
            ]
        ],
        colWidths=[78 * mm, 106 * mm],
    )
    header.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    return [header]


def _pdf_title(
    *,
    period_name: str,
    village_name: str | None,
    styles: dict[str, ParagraphStyle],
) -> list[Any]:
    title = (
        "PHIẾU BÁO CÁO SỐ LIỆU VĂN HÓA – XÃ HỘI ĐỊNH KỲ"
        if village_name
        else "BÁO CÁO TỔNG HỢP SỐ LIỆU VĂN HÓA – XÃ HỘI"
    )
    return [
        _pdf_paragraph(title, styles["title"]),
        _pdf_paragraph(f"Kỳ báo cáo: {period_name}", styles["header"]),
        Spacer(1, 3 * mm),
        _pdf_paragraph(
            f"Phạm vi: {village_name or 'Toàn xã'}  |  "
            f"Thời điểm kết xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}",
            styles["small"],
        ),
        Spacer(1, 3 * mm),
    ]


def _pdf_signature(styles: dict[str, ParagraphStyle], width: float) -> Table:
    table = Table(
        [
            [
                _pdf_paragraph(
                    "NGƯỜI LẬP BIỂU\n(Ký, ghi rõ họ tên)\n\n\n",
                    styles["header"],
                ),
                _pdf_paragraph(
                    "TM. ỦY BAN NHÂN DÂN XÃ\n(Ký số hoặc ký, đóng dấu)\n\n\n",
                    styles["header"],
                ),
            ]
        ],
        colWidths=[width / 2, width / 2],
    )
    table.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP")]))
    return table


def _pdf_village_story(
    *,
    report: dict[str, Any] | None,
    village_name: str,
    styles: dict[str, ParagraphStyle],
    usable_width: float,
) -> list[Any]:
    report = report or {}
    story: list[Any] = []
    metadata = [
        ["Thông tin", "Nội dung"],
        ["Đơn vị báo cáo", village_name],
        ["Người lập báo cáo", report.get("reporter_name") or ""],
        ["Chức danh", report.get("reporter_title") or "Cán bộ thôn"],
        ["Số điện thoại", report.get("reporter_phone") or ""],
        ["Thời điểm nộp", _format_submitted_at(report.get("submitted_at"))],
        ["Trạng thái", _submission_status_label(report) if report else "Chưa nộp"],
    ]
    story.extend(
        [
            _pdf_table(
                metadata,
                styles=styles,
                column_widths=[48 * mm, usable_width - 48 * mm],
                first_column_left=True,
            ),
            Spacer(1, 3 * mm),
            _pdf_paragraph("I. SỐ LIỆU BÁO CÁO", styles["heading"]),
        ]
    )
    values = report.get("values") if isinstance(report.get("values"), dict) else {}
    rows = [["Mã CT", "Tên chỉ tiêu", "Đơn vị tính", "Số liệu", "Ghi chú"]]
    rows.extend(
        [code, name, unit, _format_value(values.get(code)), ""]
        for code, (name, unit) in INDICATORS_DICT.items()
    )
    story.append(
        _pdf_table(
            rows,
            styles=styles,
            column_widths=[
                18 * mm,
                usable_width - 18 * mm - 28 * mm - 24 * mm - 30 * mm,
                28 * mm,
                24 * mm,
                30 * mm,
            ],
            first_column_left=True,
        )
    )
    story.append(
        KeepTogether(
            [
                _pdf_paragraph("II. XÁC NHẬN", styles["heading"]),
                _pdf_paragraph(
                    "Người lập biểu chịu trách nhiệm rà soát số liệu, đối chiếu hồ sơ "
                    "nguồn và xác nhận nội dung trước khi gửi phê duyệt.",
                    styles["body"],
                ),
                Spacer(1, 4 * mm),
                _pdf_signature(styles, usable_width),
            ]
        )
    )
    return story


def _pdf_summary_story(
    *,
    period_name: str,
    reports_data: list[dict[str, Any]],
    villages_map: dict[str, str],
    period_id: str | None,
    styles: dict[str, ParagraphStyle],
    usable_width: float,
) -> list[Any]:
    report_by_village = _report_index(reports_data)
    villages = _ordered_village_items(villages_map)
    on_time = sum(
        1 for report in reports_data if report.get("timeliness_status") == "on_time"
    )
    late = sum(
        1 for report in reports_data if report.get("timeliness_status") == "late"
    )
    story: list[Any] = [
        _pdf_paragraph("I. TÌNH HÌNH NỘP BÁO CÁO", styles["heading"]),
        _pdf_table(
            [
                ["Tổng số thôn", "Đã nộp", "Đúng hạn", "Trễ hạn", "Chưa nộp", "Tỷ lệ nộp"],
                [
                    len(villages),
                    len(reports_data),
                    on_time,
                    late,
                    max(0, len(villages) - len(reports_data)),
                    f"{(len(reports_data) * 100 / len(villages)):.1f}%"
                    if villages
                    else "0%",
                ],
            ],
            styles=styles,
            column_widths=[usable_width / 6] * 6,
        ),
        Spacer(1, 3 * mm),
    ]
    progress = [["STT", "Thôn", "Thời điểm nộp", "Trạng thái"]]
    for index, (village_id, village_name) in enumerate(villages, start=1):
        report = report_by_village.get(str(village_id))
        progress.append(
            [
                index,
                village_name,
                _format_submitted_at(report.get("submitted_at")) if report else "",
                _submission_status_label(report) if report else "Chưa nộp",
            ]
        )
    story.extend(
        [
            _pdf_table(
                progress,
                styles=styles,
                column_widths=[15 * mm, 75 * mm, 48 * mm, usable_width - 138 * mm],
                first_column_left=True,
            ),
            Spacer(1, 3 * mm),
            _pdf_paragraph(
                "II. CHỈ SỐ NGỮ NGHĨA — HỒ SƠ ĐÃ DUYỆT/KHÓA",
                styles["heading"],
            ),
            _pdf_table(
                _semantic_summary_rows(
                    period_name=period_name,
                    reports_data=reports_data,
                    villages_map=villages_map,
                    period_id=period_id,
                ),
                styles=styles,
                column_widths=[
                    48 * mm,
                    22 * mm,
                    28 * mm,
                    32 * mm,
                    usable_width - 130 * mm,
                ],
                first_column_left=True,
            ),
            PageBreak(),
            _pdf_paragraph(
                "III. BẢNG TỔNG HỢP CHỈ TIÊU THEO THÔN",
                styles["heading"],
            ),
        ]
    )
    totals, incomplete = _submitted_totals(reports_data)
    for group_name, codes in INDICATOR_GROUPS:
        rows: list[list[Any]] = [["STT", "Thôn", *codes]]
        for index, (village_id, village_name) in enumerate(villages, start=1):
            report = report_by_village.get(str(village_id))
            values = (
                report.get("values")
                if report and isinstance(report.get("values"), dict)
                else {}
            )
            rows.append(
                [
                    index,
                    village_name,
                    *[_format_value(values.get(code)) for code in codes],
                ]
            )
        rows.append(
            [
                "",
                "TỔNG CỘNG",
                *[
                    "—" if code in incomplete else _format_value(totals[code])
                    for code in codes
                ],
            ]
        )
        indicator_width = (usable_width - 15 * mm - 72 * mm) / len(codes)
        if group_name == INDICATOR_GROUPS[2][0]:
            story.append(PageBreak())
        story.extend(
            [
                _pdf_paragraph(group_name, styles["subheading"]),
                _pdf_table(
                    rows,
                    styles=styles,
                    column_widths=[
                        15 * mm,
                        72 * mm,
                        *[indicator_width for _ in codes],
                    ],
                    first_column_left=True,
                ),
                Spacer(1, 3 * mm),
            ]
        )
    dictionary = [["Mã CT", "Tên chỉ tiêu", "Đơn vị tính"]]
    dictionary.extend(
        [code, name, unit] for code, (name, unit) in INDICATORS_DICT.items()
    )
    story.extend(
        [
            PageBreak(),
            _pdf_paragraph("IV. DANH MỤC CHỈ TIÊU", styles["heading"]),
            _pdf_table(
                dictionary,
                styles=styles,
                column_widths=[24 * mm, usable_width - 64 * mm, 40 * mm],
                first_column_left=True,
            ),
            PageBreak(),
            _pdf_paragraph(
                "V. NGUỒN DỮ LIỆU VÀ CẢNH BÁO",
                styles["heading"],
            ),
        ]
    )
    sources = [["Thôn", "Nguồn nhập", "Trạng thái", "Đúng hạn", "Thời điểm nộp"]]
    for village_id, village_name in villages:
        report = report_by_village.get(str(village_id))
        sources.append(
            [
                village_name,
                _source_label(report.get("report_source")) if report else "",
                _workflow_label(report.get("workflow_status")) if report else "Chưa nộp",
                _timeliness_label(report.get("timeliness_status")) if report else "",
                _format_submitted_at(report.get("submitted_at")) if report else "",
            ]
        )
    story.append(
        _pdf_table(
            sources,
            styles=styles,
            column_widths=[58 * mm, 38 * mm, 42 * mm, 36 * mm, usable_width - 174 * mm],
            first_column_left=True,
        )
    )
    warnings = [["Thôn", "Chỉ tiêu", "Loại cảnh báo", "Nội dung", "Trạng thái"]]
    for report in reports_data:
        village_name = villages_map.get(
            str(report.get("village_id")), str(report.get("village_id") or "")
        )
        for flag in report.get("validation_flags") or []:
            warnings.append(
                [
                    village_name,
                    flag.get("ct_code") or "",
                    flag.get("error_type") or "",
                    flag.get("message") or "",
                    "Đã xử lý" if flag.get("resolved") else "Cần xem",
                ]
            )
    if len(warnings) == 1:
        warnings.append(
            ["", "", "", "Không có cảnh báo dữ liệu trong phạm vi xuất báo cáo.", ""]
        )
    story.extend(
        [
            Spacer(1, 4 * mm),
            _pdf_table(
                warnings,
                styles=styles,
                column_widths=[48 * mm, 25 * mm, 38 * mm, usable_width - 151 * mm, 40 * mm],
                first_column_left=True,
            ),
            Spacer(1, 6 * mm),
            _pdf_signature(styles, usable_width),
        ]
    )
    return story


def generate_pdf_file(
    period_name: str,
    reports_data: list,
    villages_map: dict,
    scope_name: str | None = None,
    *,
    period_id: str | None = None,
) -> bytes:
    """Generate a Unicode PDF mirroring the administrative DOCX structure."""
    styles = _pdf_styles()
    page_size = A4 if scope_name else landscape(A4)
    buffer = io.BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=page_size,
        leftMargin=30 * mm,
        rightMargin=18 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
        title=(
            "Phiếu báo cáo số liệu văn hóa – xã hội"
            if scope_name
            else "Báo cáo tổng hợp số liệu văn hóa – xã hội xã Bà Nà"
        ),
        author="Ủy ban nhân dân xã Bà Nà",
        subject=f"Kỳ báo cáo: {_safe_text(period_name)}",
    )
    story: list[Any] = [
        *_pdf_administrative_header(styles=styles, village_name=scope_name),
        *_pdf_title(
            period_name=period_name,
            village_name=scope_name,
            styles=styles,
        ),
    ]
    usable_width = page_size[0] - 48 * mm
    if scope_name:
        story.extend(
            _pdf_village_story(
                report=reports_data[0] if reports_data else None,
                village_name=scope_name,
                styles=styles,
                usable_width=usable_width,
            )
        )
    else:
        story.extend(
            _pdf_summary_story(
                period_name=period_name,
                reports_data=reports_data,
                villages_map=villages_map,
                period_id=period_id,
                styles=styles,
                usable_width=usable_width,
            )
        )

    def add_page_number(canvas: Any, doc_template: Any) -> None:
        canvas.saveState()
        canvas.setFont("BaNaTimes", 10)
        canvas.drawCentredString(page_size[0] / 2, 10 * mm, f"Trang {doc_template.page}")
        canvas.restoreState()

    document.build(
        story,
        onFirstPage=add_page_number,
        onLaterPages=add_page_number,
    )
    return buffer.getvalue()
